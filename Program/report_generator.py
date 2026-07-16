"""
Report generator for creating professional analysis reports
"""

from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field
from html import escape
from html.parser import HTMLParser
import importlib.util
import numpy as np
from datetime import datetime
import base64
import io
import os
import re
from data_loader import GrainSizeData
from analysis.comparison_snapshot import (
    ComparisonSnapshotOptions,
    DatasetAnalysisInput,
    build_comparison_snapshot,
)
from k_calculations import KCalculationResult
from k_aggregation import (
    KAggregationOptions,
    UNGROUPED_LABEL,
    build_k_result_summary,
    k_scope_value_series,
)
from grain_classification import (
    ISO14688,
    cu_label as _gc_cu_label,
    permeability_class as _gc_perm_class,
    cc_label as _gc_cc_label,
    sedimentology_descriptor as _gc_sed_descriptor,
)
from calculation_internals import compute_calculation_internals
from gui.plot_constants import classify_k_status
from gui.plot_context import (
    REPORT_EXPORT_PLOT_K_UNIT,
    context_with_style,
    convert_k_to_display,
    k_axis_label_for_unit,
    plot_context_value,
)
from gui.report_export_plot_colors import k_scope_plot_colors


def _get_plot_export():
    """Lazy import to avoid circular dependency (plot_export -> gui -> report_generator)."""
    import plot_export as _pe
    return _pe
from report_tables import analyze_report_tables, externalize_report_tables, generate_excel_appendix
from report_model import (
    AppendixLabelConfig,
    HeadingBlock,
    HtmlBlock,
    ImageBlock,
    PageBreakBlock,
    ParagraphBlock,
    ReportDocument,
    ReportSection,
    TableBlock,
)


DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
REPORT_CLASS_FRACTION_HEATMAP_MIN_UNITS = 12
REPORT_LARGE_BATCH_MIN_SAMPLES = 12
REPORT_LANDSCAPE_FIGSIZE = (13, 7.4)


class ReportCancelled(Exception):
    """Raised inside a generator when the caller's cancel_check returns True."""


@dataclass
class _HtmlNode:
    tag: str
    attrs: dict[str, str] = field(default_factory=dict)
    children: list[Any] = field(default_factory=list)


class _HtmlTreeBuilder(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.root = _HtmlNode("document")
        self._stack = [self.root]

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        node = _HtmlNode(tag, {key: value or "" for key, value in attrs})
        self._stack[-1].children.append(node)
        self._stack.append(node)

    def handle_startendtag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        node = _HtmlNode(tag, {key: value or "" for key, value in attrs})
        self._stack[-1].children.append(node)

    def handle_endtag(self, tag: str) -> None:
        while len(self._stack) > 1:
            node = self._stack.pop()
            if node.tag == tag:
                break

    def handle_data(self, data: str) -> None:
        if data:
            self._stack[-1].children.append(data)


class ReportGenerator:
    """
    Generates professional reports for grain size analysis and K-value calculations
    """

    def __init__(self):
        self._scheme = ISO14688  # Active classification scheme; set via set_scheme()

        # Brand-aware professional report stylesheet.
        # --brand is the single color token; _get_branded_style() overrides it.
        self.report_style = """
        <style>
            :root {
                --brand:       #2c3e50;
                --brand-light: rgba(44,62,80,0.08);
                --text:        #1a1a1a;
                --text-mid:    #444444;
                --text-muted:  #6c757d;
                --border:      #d0d0d0;
                --bg:          #ffffff;
                --bg-alt:      #f7f7f7;
            }

            * { box-sizing: border-box; margin: 0; padding: 0; }

            @page {
                size: A4;
                margin: 20mm 20mm 25mm 20mm;
            }

            @page comparison-landscape {
                size: A4 landscape;
                margin: 14mm 16mm 18mm 16mm;
            }

            @media print {
                body {
                    width: auto !important;
                    max-width: none !important;
                    margin: 0 !important;
                    padding: 0 !important;
                }
                .page-break { page-break-before: always; }
                .no-break   { page-break-inside: avoid; }
                h1, h2, h3  { page-break-after: avoid; }
                table       { page-break-inside: auto; break-inside: auto; }
                thead       { display: table-header-group; }
                tfoot       { display: table-footer-group; }
                tr          { page-break-inside: avoid; break-inside: avoid; }
                .page-header { display: flex !important; }
                .report-top-bar { display: none; }
                .plot-container {
                    margin: 10px 0 !important;
                    padding: 8px !important;
                }
                .landscape-plot-page {
                    page: comparison-landscape;
                    break-before: page;
                    break-after: page;
                    width: auto;
                    margin: 0;
                    padding: 0;
                }
                .landscape-plot-page h2 {
                    margin-top: 0;
                    font-size: 12pt;
                }
                .landscape-plot-page .plot-container img {
                    width: 100%;
                    max-height: 150mm;
                    object-fit: contain;
                }
                .footer {
                    margin-top: 0 !important;
                    padding-top: 6px !important;
                }
            }

            body {
                font-family: 'Calibri', 'Georgia', serif;
                line-height: 1.55;
                color: var(--text);
                background: var(--bg);
                max-width: 820px;
                margin: 0 auto;
                padding: 0 50px 60px 50px;
                font-size: 10.5pt;
            }

            /* ── Branded top bar — bleeds to body edges ──────── */
            .report-top-bar {
                height: 6px;
                background: var(--brand);
                margin: 0 -50px 40px -50px;
            }

            /* ── Cover page — bleeds to body edges ───────────── */
            .cover-page {
                padding: 50px 0 48px;
                min-height: 560px;
                display: flex;
                flex-direction: column;
                border-bottom: 1px solid var(--border);
                margin-bottom: 40px;
            }

            .cover-brand-block {
                margin-bottom: 64px;
            }

            .cover-title {
                font-size: 34px;
                font-weight: 700;
                color: var(--brand);
                margin-bottom: 10px;
                letter-spacing: -0.5px;
                line-height: 1.15;
            }

            .cover-subtitle {
                font-size: 15px;
                color: var(--text-mid);
                font-weight: 400;
                margin-bottom: 0;
            }

            .cover-meta {
                margin-top: auto;
                font-size: 10pt;
                color: var(--text-mid);
                line-height: 2;
                border-top: 1px solid var(--border);
                padding-top: 18px;
            }

            /* ── Typography ──────────────────────────────────── */
            h1 {
                font-size: 20pt;
                font-weight: 700;
                color: var(--text);
                margin: 0 0 8px 0;
                padding-bottom: 10px;
                border-bottom: 3px solid var(--brand);
            }

            h2 {
                font-size: 13pt;
                font-weight: 700;
                color: var(--brand);
                margin: 32px 0 12px 0;
                padding-bottom: 5px;
                border-bottom: 1px solid var(--border);
            }

            h3 {
                font-size: 11pt;
                font-weight: 700;
                color: var(--text);
                margin: 20px 0 8px 0;
            }

            h4 {
                font-size: 10.5pt;
                font-weight: 600;
                color: var(--text-mid);
                margin: 14px 0 6px 0;
            }

            p { margin: 8px 0; line-height: 1.55; }

            /* ── Stat cards ───────────────────────────────────── */
            .summary-stats {
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
                gap: 12px;
                margin: 20px 0;
            }

            .stat-card {
                background: var(--bg);
                padding: 14px 12px 10px;
                border: 1px solid var(--border);
                border-top: 3px solid var(--brand);
                text-align: center;
            }

            .stat-label {
                font-size: 8.5pt;
                color: var(--text-muted);
                font-weight: 600;
                text-transform: uppercase;
                letter-spacing: 0.5px;
                margin-bottom: 6px;
            }

            .stat-value {
                font-size: 19pt;
                font-weight: 700;
                color: var(--brand);
                line-height: 1.1;
            }

            .stat-unit {
                font-size: 8.5pt;
                color: var(--text-muted);
                margin-top: 3px;
            }

            /* ── Tables ───────────────────────────────────────── */
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 14px 0 18px 0;
                font-size: 9.25pt;
                line-height: 1.35;
                background: var(--bg);
                border-top: 2px solid var(--brand);
                border-bottom: 2px solid var(--brand);
                font-variant-numeric: tabular-nums;
            }

            thead { background: transparent; }

            th {
                color: var(--text);
                background: var(--brand-light);
                padding: 7px 9px;
                text-align: left;
                vertical-align: bottom;
                font-weight: 700;
                font-size: 8.75pt;
                border-bottom: 1.5px solid var(--brand);
            }

            td {
                padding: 6px 9px;
                border-bottom: 1px solid var(--border);
                vertical-align: top;
                color: var(--text);
            }

            tbody tr:last-child td { border-bottom: none; }
            tbody tr.table-group-start td { border-top: 1.5px solid var(--brand); }
            td.num, th.num, td.text-right, th.text-right { text-align: right; }
            td.center, th.center, td.text-center, th.text-center { text-align: center; }
            td.note { color: var(--text-mid); font-size: 8.5pt; }
            .table-empty { text-align: center; color: var(--text-muted); font-style: italic; }

            .table-compact th { padding: 5px 7px; font-size: 8.25pt; }
            .table-compact td { padding: 4px 7px; }
            .table-wide {
                table-layout: fixed;
                font-size: 7.75pt;
            }
            .table-wide th,
            .table-wide td {
                padding: 4px 5px;
                overflow-wrap: anywhere;
                word-break: normal;
            }
            .table-wide th:first-child,
            .table-wide td:first-child { width: 16%; }
            .table-wide th:last-child,
            .table-wide td:last-child { width: 16%; }
            .table-pair {
                display: grid;
                grid-template-columns: 1fr;
                gap: 12px;
                margin: 10px 0 18px 0;
            }
            .table-pair table { margin: 0; }
            .table-caption {
                color: var(--text-muted);
                font-size: 8.25pt;
                margin: -10px 0 16px 0;
            }

            .metadata {
                background: var(--bg-alt);
                border: 1px solid var(--border);
                border-left: 3px solid var(--brand);
                padding: 14px 16px;
                margin: 18px 0;
                font-size: 9.5pt;
            }

            .metadata-grid {
                display: grid;
                grid-template-columns: 140px 1fr;
                gap: 6px 12px;
                line-height: 1.6;
            }

            .metadata-label { font-weight: 700; color: var(--text-mid); }
            .metadata-value { color: var(--text); }

            /* ── Info / status boxes ──────────────────────────── */
            .info-box {
                background: var(--bg-alt);
                border: 1px solid var(--border);
                border-left: 3px solid var(--brand);
                padding: 12px 14px;
                margin: 12px 0;
            }

            .info-box h3 { margin-top: 0; font-size: 10pt; color: var(--brand); }
            .info-box p  { margin: 4px 0; }

            .warning-box {
                background: #fffbf0;
                border: 1px solid #e6c200;
                border-left: 3px solid #e6a200;
                padding: 12px 14px;
                margin: 12px 0;
            }

            .success-box {
                background: #f4f9f4;
                border: 1px solid #85c285;
                border-left: 3px solid #4caf50;
                padding: 12px 14px;
                margin: 12px 0;
            }

            .error-box {
                background: #fff4f4;
                border: 1px solid #e08080;
                border-left: 3px solid #cc3333;
                padding: 12px 14px;
                margin: 12px 0;
            }

            /* ── Plots ────────────────────────────────────────── */
            .plot-container {
                text-align: center;
                margin: 20px 0;
                padding: 12px;
                background: var(--bg);
                border: 1px solid var(--border);
            }

            .plot-container img { max-width: 100%; height: auto; }

            .comparison-plot-page { break-inside: avoid; }

            .figure-caption {
                font-size: 9pt;
                color: var(--text-muted);
                font-style: italic;
                margin-top: 8px;
            }

            /* ── Appendix ─────────────────────────────────────── */
            .appendix-section {
                margin-top: 48px;
                padding-top: 24px;
                border-top: 2px solid var(--brand);
            }

            .appendix-title {
                font-size: 20px;
                font-weight: 700;
                color: var(--brand);
                margin-bottom: 12px;
                text-transform: uppercase;
                letter-spacing: 1px;
            }

            .appendix-item {
                margin: 24px 0;
                padding: 16px;
                background: var(--bg-alt);
                border-left: 3px solid var(--brand);
            }

            .appendix-item-title {
                font-size: 10pt;
                font-weight: 700;
                color: var(--brand);
                margin-bottom: 10px;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }

            .appendix-subsection {
                margin-top: 18px;
            }

            .appendix-subtitle {
                margin: 0 0 10px 0;
                color: var(--text);
                font-size: 12pt;
                font-weight: 700;
            }

            /* ── Footer ───────────────────────────────────────── */
            .footer {
                margin-top: 56px;
                padding-top: 12px;
                border-top: 2px solid var(--brand);
                display: flex;
                justify-content: space-between;
                color: var(--text-muted);
                font-size: 8.5pt;
                line-height: 1.6;
            }

            /* ── Page header (print only) ─────────────────────── */
            .page-header {
                display: none;
                font-size: 8.5pt;
                color: var(--text-muted);
                padding-bottom: 8px;
                border-bottom: 1px solid var(--border);
                margin-bottom: 16px;
                justify-content: space-between;
            }

            /* ── Dividers ─────────────────────────────────────── */
            .section-divider, hr {
                border: none;
                height: 1px;
                background: var(--border);
                margin: 28px 0;
            }

            /* ── Lists ────────────────────────────────────────── */
            ul, ol { margin: 10px 0; padding-left: 22px; }
            li { margin: 4px 0; line-height: 1.55; }

            /* ── Badges ───────────────────────────────────────── */
            .badge {
                display: inline-block;
                padding: 2px 8px;
                font-size: 8.5pt;
                font-weight: 600;
                border-radius: 2px;
                background: var(--brand-light);
                color: var(--brand);
                border: 1px solid var(--brand);
            }

            .badge-success  { background: #edf7ed; color: #2a7a2a; border-color: #4caf50; }
            .badge-warning  { background: #fffbf0; color: #8a6200; border-color: #e6a200; }
            .badge-danger   { background: #fff4f4; color: #aa2222; border-color: #cc3333; }
            .badge-info     { background: var(--brand-light); color: var(--brand); border-color: var(--brand); }
            .badge-secondary{ background: #f0f0f0; color: #666;   border-color: #ccc; }

            /* ── Utilities ────────────────────────────────────── */
            .text-center { text-align: center; }
            .text-right  { text-align: right; }
            .text-muted  { color: var(--text-muted); }
            .highlight   { background: #fff8cc; padding: 1px 4px; }
            strong, b    { font-weight: 700; }

            code {
                font-family: 'Courier New', monospace;
                background: var(--bg-alt);
                padding: 1px 5px;
                border-radius: 2px;
                font-size: 9pt;
            }
        </style>
        """

    def set_scheme(self, scheme) -> None:
        """Set the active classification scheme used for all generated reports."""
        self._scheme = scheme

    def _get_branded_style(self, brand=None) -> str:
        """Return report CSS with --brand CSS variable set to the brand color."""
        if brand is None:
            return self.report_style
        color = brand.primary_color
        # Compute a transparent tint for --brand-light
        try:
            r = int(color[1:3], 16)
            g = int(color[3:5], 16)
            b = int(color[5:7], 16)
            light = f"rgba({r},{g},{b},0.08)"
        except (ValueError, IndexError):
            light = "rgba(44,62,80,0.08)"
        return (
            self.report_style
            .replace("--brand:       #2c3e50;", f"--brand:       {color};")
            .replace("--brand-light: rgba(44,62,80,0.08);", f"--brand-light: {light};")
        )

    @staticmethod
    def _esc(value: Any) -> str:
        return escape("" if value is None else str(value), quote=True)

    @staticmethod
    def _effective_porosity_value(dataset: GrainSizeData) -> Optional[float]:
        if hasattr(dataset, "effective_porosity"):
            return dataset.effective_porosity()
        current = getattr(dataset, "current_porosity", None)
        if current is not None:
            return current
        calculated = getattr(dataset, "calculated_porosity", None)
        if calculated is not None:
            return calculated
        return getattr(dataset, "porosity", None)

    def _porosity_display(self, dataset: GrainSizeData, value: Optional[float] = None) -> str:
        porosity = self._effective_porosity_value(dataset) if value is None else value
        if porosity is None:
            return "N/A"
        try:
            return f"{float(porosity):.4f}"
        except (TypeError, ValueError):
            return self._esc(porosity)

    def _porosity_source_display(self, dataset: GrainSizeData) -> str:
        if hasattr(dataset, "porosity_source_label"):
            return dataset.porosity_source_label()
        return "Current dataset value"

    def _note_html(self, value: Any) -> str:
        return self._esc(value).replace("\n", "<br>")

    @staticmethod
    def _extract_body_contents(html_text: str) -> str:
        match = re.search(r"<body[^>]*>(.*)</body>", html_text, flags=re.IGNORECASE | re.DOTALL)
        return match.group(1).strip() if match else html_text.strip()

    @staticmethod
    def _strip_leading_report_markup(body_html: str) -> str:
        body_html = re.sub(r"^\s*<div class=\"report-top-bar\"></div>\s*", "", body_html, count=1, flags=re.DOTALL)
        body_html = re.sub(
            r"^\s*<h1>\s*Hydraulic Conductivity Analysis Report\s*</h1>\s*",
            "",
            body_html,
            count=1,
            flags=re.IGNORECASE | re.DOTALL,
        )
        return body_html.strip()

    @staticmethod
    def _summarize_sample_field(sample_details: List[Dict[str, Any]], key: str, suffix: str = "") -> str:
        values = [item.get(key) for item in sample_details if item.get(key) is not None]
        if not values:
            return "N/A"
        try:
            numeric = [float(value) for value in values]
        except (TypeError, ValueError):
            unique = []
            for value in values:
                if value not in unique:
                    unique.append(value)
            if len(unique) == 1:
                return f"{unique[0]}{suffix}"
            return "Varies by sample"

        low = min(numeric)
        high = max(numeric)
        low_text = f"{low:.3f}".rstrip("0").rstrip(".")
        high_text = f"{high:.3f}".rstrip("0").rstrip(".")
        if abs(high - low) < 1e-9:
            return f"{low_text}{suffix}"
        return f"Varies by sample ({low_text} to {high_text}{suffix})"

    @staticmethod
    def _coerce_appendix_label_config(config: Optional[Any]) -> AppendixLabelConfig:
        if config is None:
            return AppendixLabelConfig()
        if isinstance(config, AppendixLabelConfig):
            return config
        if isinstance(config, dict):
            return AppendixLabelConfig(
                mode=config.get("mode", "auto"),
                scheme=config.get("scheme", "alpha"),
                layout=config.get("layout", "separate"),
                prefix=config.get("prefix", "Appendix "),
                alpha_numeric_root=config.get("alpha_numeric_root", "A"),
                single_label=config.get("single_label", ""),
                manual_labels=dict(config.get("manual_labels") or {}),
            )
        raise TypeError("appendix_label_config must be an AppendixLabelConfig or dict")

    @staticmethod
    def docx_export_available() -> bool:
        return DOCX_AVAILABLE

    @staticmethod
    def analyze_tables(html_text: str):
        return analyze_report_tables(html_text)

    @staticmethod
    def generate_excel_appendix(
        html_text: str,
        title: str = "Report data appendix",
        accent_color: str = "#2c3e50",
    ) -> bytes:
        return generate_excel_appendix(
            html_text,
            title=title,
            accent_color=accent_color,
        )

    @staticmethod
    def externalize_report_tables(html_text: str, table_titles: dict[str, str]) -> str:
        return externalize_report_tables(html_text, table_titles)

    @staticmethod
    def _node_classes(node: _HtmlNode) -> set[str]:
        return {name for name in node.attrs.get("class", "").split() if name}

    @staticmethod
    def _child_nodes(node: _HtmlNode) -> list[_HtmlNode]:
        return [child for child in node.children if isinstance(child, _HtmlNode)]

    @classmethod
    def _node_text(cls, node: _HtmlNode) -> str:
        parts: list[str] = []
        for child in node.children:
            if isinstance(child, str):
                parts.append(child)
            else:
                parts.append(cls._node_text(child))
        return re.sub(r"\s+", " ", " ".join(parts)).strip()

    @classmethod
    def _find_first_node(cls, node: _HtmlNode, tag: str) -> Optional[_HtmlNode]:
        if node.tag.lower() == tag.lower():
            return node
        for child in cls._child_nodes(node):
            match = cls._find_first_node(child, tag)
            if match is not None:
                return match
        return None

    @staticmethod
    def _parse_html_tree(html_text: str) -> _HtmlNode:
        parser = _HtmlTreeBuilder()
        parser.feed(html_text)
        parser.close()
        return parser.root

    @staticmethod
    def _hex_to_rgb_triplet(hex_color: str) -> tuple[int, int, int]:
        color = (hex_color or "#2c3e50").strip().lstrip("#")
        if len(color) == 3:
            color = "".join(ch * 2 for ch in color)
        if len(color) != 6:
            return (44, 62, 80)
        return tuple(int(color[i:i + 2], 16) for i in range(0, 6, 2))

    def _set_docx_heading_color(self, paragraph, brand_rgb: tuple[int, int, int], ctx: dict[str, Any]) -> None:
        rgb = ctx["RGBColor"](*brand_rgb)
        for run in paragraph.runs:
            run.font.color.rgb = rgb

    def _style_docx_heading(self, paragraph, level: int,
                            brand_rgb: tuple[int, int, int], ctx: dict[str, Any]) -> None:
        self._set_docx_heading_color(paragraph, brand_rgb, ctx)
        sizes = {0: 22, 1: 18, 2: 13, 3: 11, 4: 10}
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = ctx["Pt"](sizes.get(level, 10))
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = ctx["Pt"](14 if level <= 2 else 9)
        paragraph.paragraph_format.space_after = ctx["Pt"](6)

    def _add_docx_heading(self, container, text: str, level: int,
                          ctx: dict[str, Any], brand_rgb: tuple[int, int, int]):
        clean = re.sub(r"\s+", " ", text).strip()
        if not clean:
            return None
        if hasattr(container, "add_heading"):
            paragraph = container.add_heading(clean, level=level)
        else:
            style = "Title" if level == 0 else f"Heading {level}"
            paragraph = container.add_paragraph(clean, style=style)
        self._style_docx_heading(paragraph, level, brand_rgb, ctx)
        return paragraph

    def _add_docx_paragraph(self, container, text: str, ctx: dict[str, Any],
                            *, bold: bool = False, align=None, style: Optional[str] = None) -> None:
        clean = re.sub(r"\s+", " ", text).strip()
        if not clean:
            return
        paragraph = container.add_paragraph(style=style)
        if align is not None:
            paragraph.alignment = align
        run = paragraph.add_run(clean)
        run.bold = bold
        run.font.name = "Calibri"
        paragraph.paragraph_format.space_after = ctx["Pt"](6)
        paragraph.paragraph_format.line_spacing = 1.15

    @staticmethod
    def _docx_rgb_hex(rgb: tuple[int, int, int]) -> str:
        return "".join(f"{value:02X}" for value in rgb)

    def _apply_docx_box_style(self, cell, fill_hex: str, accent_hex: str,
                              ctx: dict[str, Any]) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(ctx["parse_xml"](
            f'<w:shd {ctx["nsdecls"]("w")} w:fill="{fill_hex}"/>'
        ))
        tc_pr.append(ctx["parse_xml"](
            f'<w:tcBorders {ctx["nsdecls"]("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="D0D0D0"/>'
            f'<w:left w:val="single" w:sz="18" w:space="0" w:color="{accent_hex}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="D0D0D0"/>'
            f'<w:right w:val="single" w:sz="6" w:space="0" w:color="D0D0D0"/>'
            f'</w:tcBorders>'
        ))
        tc_pr.append(ctx["parse_xml"](
            f'<w:tcMar {ctx["nsdecls"]("w")}>'
            f'<w:top w:w="140" w:type="dxa"/>'
            f'<w:left w:w="180" w:type="dxa"/>'
            f'<w:bottom w:w="140" w:type="dxa"/>'
            f'<w:right w:w="180" w:type="dxa"/>'
            f'</w:tcMar>'
        ))

    def _render_docx_box(self, container, node: _HtmlNode, ctx: dict[str, Any],
                         brand_rgb: tuple[int, int, int], state: dict[str, bool],
                         fill_hex: str) -> None:
        table = container.add_table(rows=1, cols=1)
        table.alignment = ctx["WD_TABLE_ALIGNMENT"].LEFT
        table.autofit = True
        cell = table.cell(0, 0)
        self._apply_docx_box_style(cell, fill_hex, self._docx_rgb_hex(brand_rgb), ctx)
        cell._tc.clear_content()
        for child in self._child_nodes(node):
            self._render_docx_node(cell, child, ctx, brand_rgb, state)
        spacer = container.add_paragraph()
        spacer.paragraph_format.space_after = ctx["Pt"](4)
        state["started_content"] = True

    def _apply_docx_header_footer(self, document, metadata: dict[str, str],
                                  brand, brand_rgb: tuple[int, int, int],
                                  ctx: dict[str, Any]) -> None:
        section = document.sections[0]
        section.header_distance = ctx["Mm"](8)
        section.footer_distance = ctx["Mm"](8)

        project = (metadata.get("project_name") or "Grain Size Analysis Report").strip()
        project_no = (metadata.get("project_no") or "").strip()
        header_parts = [project]
        if project_no:
            header_parts.append(project_no)
        org_name = getattr(brand, "org_name", "") if brand is not None else ""
        if org_name and org_name not in header_parts:
            header_parts.append(org_name)

        header = section.header.paragraphs[0]
        header.text = ""
        header_run = header.add_run(" | ".join(header_parts))
        header_run.bold = True
        header_run.font.name = "Calibri"
        header_run.font.size = ctx["Pt"](8.5)
        header_run.font.color.rgb = ctx["RGBColor"](*brand_rgb)
        header.paragraph_format.space_after = ctx["Pt"](2)

        footer_bits = []
        client = (metadata.get("client") or "").strip()
        analyst = (metadata.get("analyst") or "").strip()
        report_date = (metadata.get("date") or "").strip()
        if client:
            footer_bits.append(client)
        if analyst:
            footer_bits.append(f"Analyst: {analyst}")
        footer_bits.append(f"Date: {report_date or datetime.now().strftime('%Y-%m-%d')}")

        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = ctx["WD_ALIGN_PARAGRAPH"].CENTER
        footer_run = footer.add_run(" | ".join(footer_bits))
        footer_run.font.name = "Calibri"
        footer_run.font.size = ctx["Pt"](8)
        footer_run.font.color.rgb = ctx["RGBColor"](108, 117, 125)

    def _render_docx_metadata_grid(self, container, node: _HtmlNode,
                                   ctx: dict[str, Any], brand_rgb: tuple[int, int, int]) -> None:
        items: list[tuple[str, str]] = []
        pending_label = ""
        for child in self._child_nodes(node):
            text = self._node_text(child)
            if not text:
                continue
            classes = self._node_classes(child)
            if "metadata-label" in classes:
                pending_label = text
            elif "metadata-value" in classes and pending_label:
                items.append((pending_label, text))
                pending_label = ""

        if not items:
            return

        table = container.add_table(rows=len(items), cols=2)
        table.style = "Table Grid"
        table.alignment = ctx["WD_TABLE_ALIGNMENT"].LEFT
        for row_index, (label, value) in enumerate(items):
            label_para = table.cell(row_index, 0).paragraphs[0]
            label_run = label_para.add_run(label.rstrip(":"))
            label_run.bold = True
            label_run.font.color.rgb = ctx["RGBColor"](*brand_rgb)
            table.cell(row_index, 1).paragraphs[0].add_run(value)

    def _render_docx_summary_stats(self, container, node: _HtmlNode, ctx: dict[str, Any]) -> None:
        cards = [
            child for child in self._child_nodes(node)
            if "stat-card" in self._node_classes(child)
        ]
        if not cards:
            return

        table = container.add_table(rows=1, cols=len(cards))
        table.style = "Table Grid"
        table.alignment = ctx["WD_TABLE_ALIGNMENT"].CENTER
        for index, card in enumerate(cards):
            label = ""
            value = ""
            unit = ""
            for child in self._child_nodes(card):
                classes = self._node_classes(child)
                text = self._node_text(child)
                if "stat-label" in classes:
                    label = text
                elif "stat-value" in classes:
                    value = text
                elif "stat-unit" in classes:
                    unit = text

            cell = table.cell(0, index)
            cell_para = cell.paragraphs[0]
            cell_para.alignment = ctx["WD_ALIGN_PARAGRAPH"].CENTER
            label_run = cell_para.add_run(label)
            label_run.bold = True
            if value:
                value_para = cell.add_paragraph()
                value_para.alignment = ctx["WD_ALIGN_PARAGRAPH"].CENTER
                value_run = value_para.add_run(value)
                value_run.bold = True
                value_run.font.size = ctx["Pt"](13)
            if unit:
                unit_para = cell.add_paragraph()
                unit_para.alignment = ctx["WD_ALIGN_PARAGRAPH"].CENTER
                unit_para.add_run(unit)

    def _iter_table_rows(self, node: _HtmlNode) -> list[_HtmlNode]:
        rows: list[_HtmlNode] = []
        for child in self._child_nodes(node):
            tag = child.tag.lower()
            if tag == "tr":
                rows.append(child)
            elif tag in {"thead", "tbody", "tfoot"}:
                rows.extend(self._iter_table_rows(child))
        return rows

    def _apply_docx_header_shading(self, cell, fill_hex: str, ctx: dict[str, Any]) -> None:
        shading = ctx["parse_xml"](f'<w:shd {ctx["nsdecls"]("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _html_span(node: _HtmlNode, name: str) -> int:
        try:
            return max(1, int(node.attrs.get(name, "1") or 1))
        except (TypeError, ValueError):
            return 1

    def _apply_docx_table_rules(self, table, fill_hex: str, ctx: dict[str, Any]) -> None:
        borders = ctx["parse_xml"](
            f'<w:tblBorders {ctx["nsdecls"]("w")}>'
            f'<w:top w:val="single" w:sz="14" w:space="0" w:color="{fill_hex}"/>'
            f'<w:left w:val="nil"/>'
            f'<w:bottom w:val="single" w:sz="14" w:space="0" w:color="{fill_hex}"/>'
            f'<w:right w:val="nil"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
            f'<w:insideV w:val="nil"/>'
            f'</w:tblBorders>'
        )
        table._tbl.tblPr.append(borders)

    def _apply_docx_cell_bottom_rule(self, cell, fill_hex: str, ctx: dict[str, Any]) -> None:
        borders = ctx["parse_xml"](
            f'<w:tcBorders {ctx["nsdecls"]("w")}>'
            f'<w:bottom w:val="single" w:sz="10" w:space="0" w:color="{fill_hex}"/>'
            f'</w:tcBorders>'
        )
        cell._tc.get_or_add_tcPr().append(borders)

    def _render_docx_table(self, container, node: _HtmlNode,
                           ctx: dict[str, Any], brand_rgb: tuple[int, int, int]) -> None:
        rows = self._iter_table_rows(node)
        if not rows:
            return

        placements = []
        occupied: set[tuple[int, int]] = set()
        col_count = 0
        for row_index, row_node in enumerate(rows):
            column_index = 0
            cells = [
                child for child in self._child_nodes(row_node)
                if child.tag.lower() in {"th", "td"}
            ]
            for cell_node in cells:
                while (row_index, column_index) in occupied:
                    column_index += 1
                colspan = self._html_span(cell_node, "colspan")
                rowspan = self._html_span(cell_node, "rowspan")
                placements.append((row_index, column_index, colspan, rowspan, cell_node))
                for occupied_row in range(row_index, row_index + rowspan):
                    for occupied_col in range(column_index, column_index + colspan):
                        occupied.add((occupied_row, occupied_col))
                column_index += colspan
                col_count = max(col_count, column_index)

        table = container.add_table(rows=len(rows), cols=max(col_count, 1))
        table.alignment = ctx["WD_TABLE_ALIGNMENT"].CENTER
        table.autofit = True
        fill_hex = "".join(f"{value:02X}" for value in brand_rgb)
        self._apply_docx_table_rules(table, fill_hex, ctx)
        font_size = ctx["Pt"](7.5 if col_count >= 7 else 9)

        header_rows: set[int] = set()
        for row_index, column_index, colspan, rowspan, cell_node in placements:
            cell = table.cell(row_index, column_index)
            if colspan > 1 or rowspan > 1:
                cell = cell.merge(
                    table.cell(row_index + rowspan - 1, column_index + colspan - 1)
                )
            paragraph = cell.paragraphs[0]
            classes = self._node_classes(cell_node)
            style_text = cell_node.attrs.get("style", "").lower()
            if "text-align: right" in style_text or "num" in classes or "text-right" in classes:
                paragraph.alignment = ctx["WD_ALIGN_PARAGRAPH"].RIGHT
            elif "text-align: center" in style_text or "center" in classes or "text-center" in classes:
                paragraph.alignment = ctx["WD_ALIGN_PARAGRAPH"].CENTER

            text = self._node_text(cell_node)
            if text:
                run = paragraph.add_run(text)
                run.font.size = font_size
                if cell_node.tag.lower() == "th":
                    run.bold = True
                    run.font.color.rgb = ctx["RGBColor"](*brand_rgb)

            if cell_node.tag.lower() == "th":
                header_rows.add(row_index)
                self._apply_docx_header_shading(cell, "EAF0F3", ctx)
                self._apply_docx_cell_bottom_rule(cell, fill_hex, ctx)

        for row_index in sorted(header_rows):
            header = ctx["parse_xml"](
                f'<w:tblHeader {ctx["nsdecls"]("w")} w:val="true"/>'
            )
            table.rows[row_index]._tr.get_or_add_trPr().append(header)

    def _render_docx_image(self, container, node: _HtmlNode, ctx: dict[str, Any]) -> None:
        src = node.attrs.get("src", "")
        if not src:
            return

        image_bytes = None
        if src.startswith("data:"):
            match = re.match(r"data:([^;]+);base64,(.*)", src, flags=re.DOTALL)
            if not match:
                return
            mime_type = match.group(1).lower()
            if "svg" in mime_type:
                return
            image_bytes = base64.b64decode(match.group(2))
        elif os.path.exists(src):
            with open(src, "rb") as fh:
                image_bytes = fh.read()

        if not image_bytes:
            return

        paragraph = container.add_paragraph()
        paragraph.alignment = ctx["WD_ALIGN_PARAGRAPH"].CENTER
        run = paragraph.add_run()
        alt_text = (node.attrs.get("alt", "") or "").lower()
        image_width = 2.0 if "logo" in alt_text else (10.2 if ctx.get("landscape") else 6.0)
        width = ctx["Inches"](image_width)
        run.add_picture(io.BytesIO(image_bytes), width=width)

    def _render_docx_node(self, container, node: _HtmlNode,
                          ctx: dict[str, Any], brand_rgb: tuple[int, int, int],
                          state: dict[str, bool]) -> None:
        tag = node.tag.lower()
        classes = self._node_classes(node)

        if tag in {"document", "html", "body"}:
            for child in self._child_nodes(node):
                self._render_docx_node(container, child, ctx, brand_rgb, state)
            return

        if tag in {"head", "style", "script", "meta", "title"}:
            return

        if tag == "div":
            if "report-top-bar" in classes or "footer" in classes:
                return
            if "page-break" in classes and state.get("started_content") and hasattr(container, "add_page_break"):
                container.add_page_break()
            box_fills = {
                "metadata": "F7F7F7",
                "info-box": "F7F7F7",
                "success-box": "F4F9F4",
                "warning-box": "FFFBF0",
                "error-box": "FFF4F4",
                "appendix-item": "F7F7F7",
            }
            for class_name, fill_hex in box_fills.items():
                if class_name in classes:
                    self._render_docx_box(container, node, ctx, brand_rgb, state, fill_hex)
                    return
            if "metadata-grid" in classes:
                self._render_docx_metadata_grid(container, node, ctx, brand_rgb)
                state["started_content"] = True
                return
            if "summary-stats" in classes:
                self._render_docx_summary_stats(container, node, ctx)
                state["started_content"] = True
                return
            if "cover-title" in classes:
                paragraph = self._add_docx_heading(
                    container, self._node_text(node), 0, ctx, brand_rgb
                )
                if paragraph is not None:
                    paragraph.alignment = ctx["WD_ALIGN_PARAGRAPH"].CENTER
                state["started_content"] = True
                return
            if "cover-subtitle" in classes:
                self._add_docx_paragraph(
                    container,
                    self._node_text(node),
                    ctx,
                    align=ctx["WD_ALIGN_PARAGRAPH"].CENTER,
                )
                state["started_content"] = True
                return

            child_nodes = self._child_nodes(node)
            if not child_nodes:
                text = self._node_text(node)
                if text:
                    self._add_docx_paragraph(container, text, ctx)
                    state["started_content"] = True
                return

            block_tags = {"div", "p", "h1", "h2", "h3", "h4", "table", "ul", "ol", "img"}
            if not any(child.tag.lower() in block_tags for child in child_nodes):
                text = self._node_text(node)
                if text:
                    self._add_docx_paragraph(container, text, ctx)
                    state["started_content"] = True
                return

            for child in child_nodes:
                self._render_docx_node(container, child, ctx, brand_rgb, state)
            return

        if tag in {"h1", "h2", "h3", "h4"}:
            heading_level = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}[tag]
            self._add_docx_heading(
                container, self._node_text(node), heading_level, ctx, brand_rgb
            )
            state["started_content"] = True
            return
        if tag == "p":
            self._add_docx_paragraph(container, self._node_text(node), ctx)
            state["started_content"] = True
            return

        if tag in {"ul", "ol"}:
            for child in self._child_nodes(node):
                self._render_docx_node(container, child, ctx, brand_rgb, state)
            return

        if tag == "li":
            self._add_docx_paragraph(container, self._node_text(node), ctx, style="List Bullet")
            state["started_content"] = True
            return

        if tag == "table":
            table_id = node.attrs.get("data-report-table", "")
            externalized_ids = ctx.get("externalized_table_ids", set())
            if table_id and table_id in externalized_ids:
                title = ctx.get("externalized_table_titles", {}).get(table_id) or table_id
                self._add_docx_paragraph(
                    container,
                    f"Large table moved to companion Excel appendix: {title}.",
                    ctx,
                    bold=True,
                )
                state["started_content"] = True
                return
            self._render_docx_table(container, node, ctx, brand_rgb)
            state["started_content"] = True
            return

        if tag == "img":
            self._render_docx_image(container, node, ctx)
            state["started_content"] = True
            return

        for child in self._child_nodes(node):
            self._render_docx_node(container, child, ctx, brand_rgb, state)

    def generate_docx_from_html(
        self,
        html_text: str,
        brand=None,
        *,
        metadata: Optional[dict[str, str]] = None,
        externalized_table_ids: Optional[set[str]] = None,
        externalized_table_titles: Optional[dict[str, str]] = None,
    ) -> bytes:
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is not installed.")

        from docx import Document
        from docx.enum.section import WD_ORIENT, WD_SECTION
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        from docx.shared import Inches, Mm, Pt, RGBColor

        document = Document()
        def configure_section(section, *, landscape: bool) -> None:
            section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
            section.page_width = Mm(297 if landscape else 210)
            section.page_height = Mm(210 if landscape else 297)
            section.top_margin = Mm(14 if landscape else 20)
            section.right_margin = Mm(16 if landscape else 20)
            section.bottom_margin = Mm(18 if landscape else 20)
            section.left_margin = Mm(16 if landscape else 20)

        configure_section(document.sections[0], landscape=False)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(10.5)

        body_html = self._extract_body_contents(html_text)
        root = self._parse_html_tree(body_html)
        body = self._find_first_node(root, "body") or root
        brand_rgb = self._hex_to_rgb_triplet(getattr(brand, "primary_color", "#2c3e50"))
        ctx = {
            "Inches": Inches,
            "Mm": Mm,
            "Pt": Pt,
            "RGBColor": RGBColor,
            "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
            "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
            "parse_xml": parse_xml,
            "nsdecls": nsdecls,
            "externalized_table_ids": set(externalized_table_ids or ()),
            "externalized_table_titles": dict(externalized_table_titles or {}),
            "landscape": False,
        }
        self._apply_docx_header_footer(document, dict(metadata or {}), brand, brand_rgb, ctx)
        state = {"started_content": False}
        landscape_active = False
        landscape_page_started = False
        for child in self._child_nodes(body):
            classes = self._node_classes(child)
            if "report-top-bar" in classes or "footer" in classes:
                continue
            wants_landscape = "landscape-plot-page" in classes
            if wants_landscape:
                if not landscape_active:
                    section = document.add_section(WD_SECTION.NEW_PAGE)
                    configure_section(section, landscape=True)
                    landscape_active = True
                    landscape_page_started = False
                elif landscape_page_started:
                    document.add_page_break()
                ctx["landscape"] = True
                self._render_docx_node(document, child, ctx, brand_rgb, state)
                landscape_page_started = True
                continue
            if landscape_active and "page-break" in classes:
                continue
            if landscape_active:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=False)
                landscape_active = False
                landscape_page_started = False
                ctx["landscape"] = False
            self._render_docx_node(document, child, ctx, brand_rgb, state)

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _render_table_block_html(self, block: TableBlock) -> str:
        html = "<table>"
        if block.columns:
            html += "<thead><tr>"
            for column in block.columns:
                html += f"<th>{self._esc(column)}</th>"
            html += "</tr></thead>"

        html += "<tbody>"
        for row in block.rows:
            html += "<tr>"
            for cell in row:
                html += f"<td>{self._esc(cell)}</td>"
            html += "</tr>"
        html += "</tbody></table>"

        if block.caption:
            html += f'<div class="figure-caption">{self._esc(block.caption)}</div>'
        return html

    def _render_report_block_html(self, block: Any) -> str:
        if isinstance(block, HeadingBlock):
            level = min(max(block.level, 1), 4)
            return f"<h{level}>{self._esc(block.text)}</h{level}>"
        if isinstance(block, ParagraphBlock):
            return f"<p>{self._note_html(block.text)}</p>"
        if isinstance(block, TableBlock):
            return self._render_table_block_html(block)
        if isinstance(block, HtmlBlock):
            return block.html
        if isinstance(block, ImageBlock):
            caption = ""
            if block.caption:
                caption = f'<div class="figure-caption">{self._esc(block.caption)}</div>'
            return (
                f'<div class="plot-container">'
                f'<img src="{self._esc(block.source)}" alt="{self._esc(block.alt_text or block.caption)}" />'
                f'{caption}'
                f'</div>'
            )
        if isinstance(block, PageBreakBlock):
            return '<div class="page-break"></div>'
        raise TypeError(f"Unsupported report block: {type(block)!r}")

    def _calculate_percentile_values(self, dataset: GrainSizeData,
                                     percentiles_list: List[int]) -> Dict[int, Optional[float]]:
        values: Dict[int, Optional[float]] = {}
        for percentile in percentiles_list:
            values[percentile] = dataset.get_percentile_size(percentile)
        return values

    def _create_percentiles_table_block(self, dataset: GrainSizeData) -> TableBlock:
        percentiles_list = [5, 10, 16, 20, 25, 30, 40, 50, 60, 75, 84, 90, 95]
        percentiles_dict = self._calculate_percentile_values(dataset, percentiles_list)
        max_val = max((value for value in percentiles_dict.values() if value is not None), default=0)
        rows: list[list[str]] = []
        for percentile in percentiles_list:
            value = percentiles_dict[percentile]
            relative = int((value / max_val) * 100) if value is not None and max_val > 0 else 0
            label = f"D{percentile}{'*' if percentile in [10, 30, 50, 60] else ''}"
            rows.append([label, f"{value:.3f}" if value is not None else "N/A", f"{relative}%"])

        return TableBlock(
            columns=["Percentile", "Size (mm)", "Relative (%)"],
            rows=rows,
        )

    def _create_data_quality_table_block(self, dataset: GrainSizeData) -> TableBlock:
        n_points = len(dataset.particle_sizes)
        size_min = min(dataset.particle_sizes)
        size_max = max(dataset.particle_sizes)
        size_range = size_max / size_min if size_min > 0 else 0

        sorted_indices = np.argsort(dataset.particle_sizes)[::-1]
        sorted_passing = [dataset.percent_passing[i] for i in sorted_indices]
        monotonic = all(sorted_passing[i] >= sorted_passing[i + 1] for i in range(len(sorted_passing) - 1))
        monotonicity_score = "Excellent" if monotonic else "Good"
        coverage_score = "Excellent" if size_range > 100 else "Good" if size_range > 10 else "Limited"
        density_score = "Excellent" if n_points > 20 else "Good" if n_points > 10 else "Adequate"
        confidence_score = "High" if (n_points > 15 and size_range > 50) else "Moderate" if n_points > 8 else "Low"

        rows = [
            ["Number of Data Points", str(n_points), density_score],
            ["Size Range", f"{size_min:.3f} - {size_max:.1f} mm", coverage_score],
            ["Span Ratio", f"{size_range:.1f}x", coverage_score],
            ["Curve Monotonicity", "Monotonic" if monotonic else "Some variation", monotonicity_score],
            ["Interpolation Confidence", "", confidence_score],
        ]
        return TableBlock(
            columns=["Quality Metric", "Value", "Assessment"],
            rows=rows,
        )

    def _create_raw_data_table_block(self, dataset: GrainSizeData) -> TableBlock:
        rows: list[list[str]] = []
        for size, passing in zip(dataset.particle_sizes, dataset.percent_passing):
            retained = 100 - passing
            rows.append([f"{size:.4f}", f"{passing:.2f}", f"{retained:.2f}"])

        return TableBlock(
            columns=["Grain Size (mm)", "Percent Passing (%)", "Percent Retained (%)"],
            rows=rows,
        )

    def _build_grain_size_appendix_document(self, dataset: GrainSizeData,
                                            sections: Dict[str, bool],
                                            appendix_label_config: Optional[Any] = None) -> ReportDocument:
        document = ReportDocument(
            title="Appendices",
            appendix_label_config=self._coerce_appendix_label_config(appendix_label_config),
        )

        if sections.get('percentiles', True):
            document.add_section(ReportSection(
                section_id="grain_percentiles",
                title="Detailed Percentile Data",
                kind="appendix",
                blocks=[HtmlBlock(self._create_percentiles_table(dataset))],
            ))
        if sections.get('data_quality', False):
            document.add_section(ReportSection(
                section_id="grain_data_quality",
                title="Data Quality Assessment",
                kind="appendix",
                blocks=[HtmlBlock(self._create_data_quality_table(dataset))],
            ))
        if sections.get('raw_data', False):
            raw_data_table = '<table class="table-compact" data-report-table="raw-measurements"><thead><tr><th>Grain Size (mm)</th><th>Percent Passing (%)</th><th>Percent Retained (%)</th></tr></thead><tbody>'
            for size, passing in zip(dataset.particle_sizes, dataset.percent_passing):
                retained = 100 - passing
                raw_data_table += f'<tr><td>{size:.4f}</td><td>{passing:.2f}</td><td>{retained:.2f}</td></tr>'
            raw_data_table += '</tbody></table>'
            document.add_section(ReportSection(
                section_id="grain_raw_data",
                title="Raw Measurement Data",
                kind="appendix",
                blocks=[HtmlBlock(raw_data_table)],
            ))

        return document

    def _render_appendix_document_html(self, document: ReportDocument) -> str:
        appendix_sections = document.appendix_sections()
        if not appendix_sections:
            return ""

        html = '<div class="appendix-section page-break">'
        html += f'<h1 class="appendix-title">{self._esc(document.title)}</h1>'

        if document.single_appendix_enabled():
            html += '<div class="appendix-item">'
            html += f'<h3 class="appendix-item-title">{self._esc(document.single_appendix_label())}</h3>'
            for section in appendix_sections:
                html += '<div class="appendix-subsection">'
                if section.title:
                    html += f'<h4 class="appendix-subtitle">{self._esc(section.title)}</h4>'
                if section.notes:
                    html += f"<p>{self._note_html(section.notes)}</p>"
                for block in section.blocks:
                    html += self._render_report_block_html(block)
                html += '</div>'
            html += '</div>'
        else:
            for section in appendix_sections:
                html += '<div class="appendix-item">'
                html += f'<h3 class="appendix-item-title">{self._esc(document.appendix_display_title(section))}</h3>'
                if section.notes:
                    html += f"<p>{self._note_html(section.notes)}</p>"
                for block in section.blocks:
                    html += self._render_report_block_html(block)
                html += '</div>'

        html += '</div>'
        return html

    def _create_cover_page(self, title: str, subtitle: str,
                           metadata: Dict[str, str], brand=None) -> str:
        """Create a professional cover page"""
        html = '<div class="cover-page page-break">'

        # Brand header (logo + org name) if branding provided
        if brand is not None:
            html += f'<div style="margin-bottom:24px;">{brand.get_logo_html(56)}</div>'
            html += (
                f'<div style="font-size:13px;font-weight:600;'
                f'color:{brand.primary_color};margin-bottom:4px;">'
                f'{self._esc(brand.org_name)}</div>'
            )
            if brand.org_subtitle:
                html += (
                    f'<div style="font-size:11px;color:#7f8c8d;'
                    f'margin-bottom:28px;">{self._esc(brand.org_subtitle)}</div>'
                )

        html += f'<div class="cover-title">{self._esc(title)}</div>'
        html += f'<div class="cover-subtitle">{self._esc(subtitle)}</div>'

        html += '<div class="cover-meta">'
        if metadata.get('project_name'):
            html += f'<div><strong>Project:</strong> {self._esc(metadata["project_name"])}</div>'
        if metadata.get('location'):
            html += f'<div><strong>Location:</strong> {self._esc(metadata["location"])}</div>'
        if metadata.get('client'):
            html += f'<div><strong>Client:</strong> {self._esc(metadata["client"])}</div>'
        if metadata.get('analyst'):
            html += f'<div><strong>Analyst:</strong> {self._esc(metadata["analyst"])}</div>'
        html += f'<div><strong>Date:</strong> {datetime.now().strftime("%B %d, %Y")}</div>'
        html += '</div>'
        html += '</div>'

        return html

    def _global_report_style(self):
        """The global report/export PlotStyle (preset + saved Customize overrides)."""
        from gui.report_plot_style import resolve_report_style
        return resolve_report_style()

    def _palette_curve_color(self) -> Optional[str]:
        """First palette colour for a single-curve plot, or None to keep the preset.

        One rule, shared with the comparison plots: a chosen colormap palette
        (Viridis/…) colours the curve; the Categorical default leaves the curve at
        the preset's colour (unchanged behaviour). The preset always governs
        typography either way.
        """
        from gui.plot_constants import CATEGORICAL_PALETTE
        from gui.report_plot_style import get_report_palette, resolve_report_palette_colors
        if get_report_palette() == CATEGORICAL_PALETTE:
            return None
        palette = resolve_report_palette_colors(1)
        return palette[0] if palette else None

    def _create_grain_size_plot(
        self,
        dataset: GrainSizeData,
        plot_context: Optional[Dict[str, Any]] = None,
        curve_color: Optional[str] = None,
    ) -> str:
        """Create grain size distribution curve and return as base64.

        The global report style is forced onto the captured context so the
        Customize panel themes this plot too (other context — unit, limits,
        text, grid/legend — is preserved). *curve_color* overrides the curve
        colour so a per-sample plot can match that sample's palette colour in the
        comparison overlay; when omitted the single-curve palette colour applies,
        so a standalone Individual report follows the palette too. Typography
        always comes from the global preset.
        """
        pe = _get_plot_export()
        style = self._global_report_style()
        if curve_color is None:
            curve_color = self._palette_curve_color()
        if curve_color:
            import dataclasses
            style = dataclasses.replace(style, curve_color=curve_color)
        return pe.export_grain_size_plot(
            dataset,
            style=style,
            plot_context=context_with_style(plot_context, style),
            show_d_lines=False,
            show_markers=False,
            classification_scheme=self._scheme,
        )

    def _create_grain_size_histogram(
        self,
        dataset: GrainSizeData,
        plot_context: Optional[Dict[str, Any]] = None,
        plot_style=None,
        bar_color: Optional[str] = None,
    ) -> str:
        """Create a class-fraction histogram for one sample."""
        style = plot_style or self._global_report_style()
        if bar_color:
            import dataclasses
            style = dataclasses.replace(style, curve_color=bar_color)
        detail = {
            "dataset": dataset,
            "k_results": [],
            "group_name": getattr(dataset, "group_name", None) or UNGROUPED_LABEL,
            "plot_context": plot_context,
        }
        spec = self._build_comparison_spec(
            [detail],
            None,
            plot_type="histogram",
            display_mode="grid",
            breakdown="dataset",
            plot_style=style,
        )
        if bar_color:
            spec.palette = [bar_color]
            spec.effective_colors = [bar_color]
            spec.color_by_name[dataset.sample_name] = bar_color
            spec.palette_authoritative = True
        spec.show_grid = bool(plot_context_value(plot_context, "show_grid", True))
        return _get_plot_export().export_comparison_spec(spec, figsize=(10, 6))

    def _create_k_value_bar_chart(
        self,
        k_results: List[KCalculationResult],
        plot_context: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Create a report K-value bar chart in the canonical plot unit.

        Style, log-K axis, grid and legend state come from captured context;
        report/export figures intentionally use m/s for cross-path parity.
        """
        valid_results = [r for r in k_results if r.k_value is not None and r.k_value > 0]
        if not valid_results:
            return ""

        unit = REPORT_EXPORT_PLOT_K_UNIT
        methods = [r.method_name for r in valid_results]
        k_values = [convert_k_to_display(r.k_value, unit) for r in valid_results]
        flagged = {r.method_name for r in valid_results if classify_k_status(r) != "OK"}
        reference_values = [
            convert_k_to_display(r.k_value, unit)
            for r in valid_results
            if classify_k_status(r) == "OK"
        ]

        pe = _get_plot_export()
        return pe.export_k_bar_chart(
            methods, k_values,
            flagged_methods=flagged,
            reference_values=reference_values,
            style=self._global_report_style(),
            show_grid=bool(plot_context_value(plot_context, "show_grid", True)),
            show_legend=bool(plot_context_value(plot_context, "show_legend", True)),
            log_y_scale=bool(plot_context_value(plot_context, "log_k_y_scale", False)),
            y_label=k_axis_label_for_unit(unit),
            title="Hydraulic Conductivity Estimates by Method",
            colors=self._k_bar_method_colors(methods),
        )

    def _k_bar_method_colors(self, methods) -> List[str]:
        """Per-method bar colours from the active palette (preset never decides).

        Colormap palettes (Viridis/…) sample one colour per method; the
        Categorical default keeps the fixed semantic method colours (Hazen=red…)
        so meaning is preserved. Either way the preset only governs typography.
        """
        from gui.plot_constants import (
            CATEGORICAL_PALETTE, METHOD_COLORS, palette_colors,
        )
        from gui.report_plot_style import get_report_palette
        palette = get_report_palette()
        if palette == CATEGORICAL_PALETTE:
            return [METHOD_COLORS.get(m, "#888888") for m in methods]
        return palette_colors(palette, len(methods))

    def _create_method_applicability_heatmap(
        self,
        k_results: List[KCalculationResult],
        plot_context: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Create method applicability status heatmap (styled from the global style)."""
        if not k_results:
            return ""
        return _get_plot_export().export_applicability_heatmap(
            k_results, style=self._global_report_style()
        )

    @staticmethod
    def _comparison_report_style(plot_style, sample_count: int):
        from gui.plot_styles import PROFESSIONAL_STYLE
        from gui.report_plot_style import get_report_style_overrides

        resolved_style = plot_style or PROFESSIONAL_STYLE
        if sample_count < REPORT_LARGE_BATCH_MIN_SAMPLES:
            return resolved_style, False
        overrides = get_report_style_overrides()
        updates = {}
        automatic_position = not (
            {"legend_loc", "legend_bbox_to_anchor"} & set(overrides)
        )
        if automatic_position:
            updates.update(
                legend_loc="upper center",
                legend_bbox_to_anchor=(0.5, -0.18),
            )
        if "legend_ncol" not in overrides:
            updates["legend_ncol"] = 0
        if updates:
            import dataclasses
            resolved_style = dataclasses.replace(resolved_style, **updates)
        return resolved_style, automatic_position

    def _build_comparison_spec(self, sample_details, comparison_snapshot, *,
                               plot_type: str, display_mode: str = "overlay",
                               breakdown: Optional[str] = None,
                               plot_style=None):
        """Capture a widget-free ComparisonPlotSpec for the report's samples.

        Comparison plots render through the same pipeline as the Comparison tab
        (group breakdown, group colours, per-dataset line styles), keyed to the
        report's own selected samples. K is shown in m/s to match the report's
        tables and K boxplot. *breakdown* forces ``"group"``/``"dataset"`` (None
        = auto: group when named groups exist).
        """
        from gui.comparison_plot_capture import build_comparison_spec
        from gui.plot_styles import PROFESSIONAL_STYLE
        from gui.plot_constants import CATEGORICAL_PALETTE
        from gui.report_plot_style import get_report_palette, resolve_report_palette_colors
        datasets = [item["dataset"] for item in sample_details]
        results_by_name = {
            item["dataset"].sample_name: list(item.get("k_results") or [])
            for item in sample_details
        }
        dataset_groups = {
            item["dataset"].sample_name: (
                item.get("group_name")
                or getattr(item["dataset"], "group_name", None)
                or "Ungrouped"
            )
            for item in sample_details
        }
        dense_layout = len(datasets) >= REPORT_LARGE_BATCH_MIN_SAMPLES
        resolved_style, automatic_legend_layout = self._comparison_report_style(
            plot_style, len(datasets)
        )
        spec = build_comparison_spec(
            datasets,
            results_by_name,
            comparison_snapshot=comparison_snapshot,
            dataset_groups=dataset_groups,
            current_plot_type=plot_type,
            display_mode=display_mode,
            breakdown=breakdown,
            style=resolved_style,
            display_unit=REPORT_EXPORT_PLOT_K_UNIT,
            classification_scheme=self._scheme,
            # Dataset/group colours follow the global report palette so every
            # comparison plot re-colours at once (Categorical → GUI defaults).
            # palette_name lets the spec re-sample per series-count so groups
            # spread across the whole colormap (not just its dark end).
            palette=resolve_report_palette_colors(max(len(datasets), 1)),
            palette_name=get_report_palette(),
            group_palette_authoritative=get_report_palette() != CATEGORICAL_PALETTE,
        )
        spec.dense_report_layout = dense_layout
        spec.automatic_report_legend_layout = automatic_legend_layout
        return spec

    def _create_comparison_grain_size_plot(self, sample_details, comparison_snapshot,
                                           breakdown: Optional[str] = None,
                                           plot_style=None) -> str:
        """Grain-size distribution comparison, matching the Comparison tab."""
        spec = self._build_comparison_spec(
            sample_details, comparison_snapshot,
            plot_type="distribution", breakdown=breakdown, plot_style=plot_style,
        )
        return _get_plot_export().export_comparison_spec(
            spec,
            figsize=REPORT_LANDSCAPE_FIGSIZE if spec.dense_report_layout else (12, 7),
        )

    def _build_report_class_fraction_spec(
        self,
        sample_details,
        comparison_snapshot,
        *,
        breakdown: Optional[str] = None,
        plot_style=None,
    ):
        """Resolve the report layout from the number of plotted units."""
        spec = self._build_comparison_spec(
            sample_details, comparison_snapshot,
            plot_type="histogram", display_mode="overlay",
            breakdown=breakdown, plot_style=plot_style,
        )
        from gui.comparison_plot_spec import histogram_units

        unit_count = len(histogram_units(spec))
        use_heatmap = unit_count >= REPORT_CLASS_FRACTION_HEATMAP_MIN_UNITS
        spec.histogram_layout = "heatmap" if use_heatmap else "bars"
        if use_heatmap:
            spec.show_legend = False
        return spec, unit_count

    def _class_fraction_layout_note(
        self,
        sample_details,
        comparison_snapshot,
        *,
        breakdown: Optional[str] = None,
        plot_style=None,
    ) -> str:
        """Explain an automatic heatmap switch in the report preview/output."""
        spec, unit_count = self._build_report_class_fraction_spec(
            sample_details,
            comparison_snapshot,
            breakdown=breakdown,
            plot_style=plot_style,
        )
        if spec.histogram_layout != "heatmap":
            return ""
        scope = "groups" if spec.use_group_breakdown else "samples"
        return (
            f"Large-batch layout: Heatmap shown for {unit_count} {scope}; "
            f"rows are {scope}, and color shows weight percent (0-100)."
        )

    def _create_comparison_grain_size_histogram(self, sample_details, comparison_snapshot,
                                                breakdown: Optional[str] = None,
                                                plot_style=None) -> str:
        """Render class fractions with an A4-friendly large-scope layout."""
        spec, unit_count = self._build_report_class_fraction_spec(
            sample_details,
            comparison_snapshot,
            breakdown=breakdown,
            plot_style=plot_style,
        )
        use_heatmap = spec.histogram_layout == "heatmap"
        figure_height = 7.4 if spec.dense_report_layout else (
            min(10.5, max(7.0, 2.8 + unit_count * 0.15))
            if use_heatmap else 7.0
        )
        return _get_plot_export().export_comparison_spec(
            spec,
            figsize=(13 if spec.dense_report_layout else 12, figure_height),
        )

    def _create_comparison_k_value_bar(self, sample_details, comparison_snapshot,
                                       breakdown: Optional[str] = None,
                                       plot_style=None) -> str:
        """Grouped K-value bar comparison (one bar series per dataset/group)."""
        spec = self._build_comparison_spec(
            sample_details, comparison_snapshot,
            plot_type="k-values", breakdown=breakdown, plot_style=plot_style,
        )
        if not spec.k_results_dict:
            return ""
        return _get_plot_export().export_comparison_spec(
            spec,
            figsize=REPORT_LANDSCAPE_FIGSIZE if spec.dense_report_layout else (12, 7),
        )

    def _create_per_sample_plots_section(self, sample_details, *,
                                         include_grain: bool, include_histogram: bool, include_kbar: bool,
                                         advance=None, colors_by_name=None) -> str:
        """Render each sample's own grain curve, class histogram and/or K-value bar.

        Lets a multi-sample (Comparison/Full) report carry per-sample detail in
        addition to the cross-sample plots. Each sample becomes a sub-block; plots
        reuse the single-sample helpers (so they honour the global report style).
        *advance* (optional ``callable(label)``) is called once per sample to drive
        a progress bar and honour cancellation (it may raise ``ReportCancelled``).
        *colors_by_name* maps a dataset's ``sample_name`` to the colour it has in
        the comparison overlay (from the active palette), so each per-sample grain
        curve matches its overlay colour while keeping the global typography.
        """
        if not (include_grain or include_histogram or include_kbar):
            return ""
        colors_by_name = colors_by_name or {}

        blocks = ""
        for item in sample_details:
            dataset = item.get("dataset")
            if dataset is None:
                continue
            label = item.get("label") or dataset.sample_name
            if advance is not None:
                advance(f"Rendering plots for {label}")
            plot_context = item.get("plot_context")
            curve_color = colors_by_name.get(dataset.sample_name)
            grain = (
                self._create_grain_size_plot(dataset, plot_context, curve_color=curve_color)
                if include_grain else ""
            )
            histogram = (
                self._create_grain_size_histogram(
                    dataset, plot_context, bar_color=curve_color
                )
                if include_histogram else ""
            )
            kbar = (
                self._create_k_value_bar_chart(list(item.get("k_results") or []), plot_context)
                if include_kbar else ""
            )
            if not (grain or histogram or kbar):
                continue
            images = ""
            if grain:
                images += (
                    '<div class="plot-container">'
                    f'<img src="{grain}" alt="{label} grain size distribution" '
                    'style="max-width: 100%; height: auto;"></div>'
                )
            if histogram:
                images += (
                    '<div class="plot-container">'
                    f'<img src="{histogram}" alt="{label} grain-size class histogram" '
                    'style="max-width: 100%; height: auto;"></div>'
                )
            if kbar:
                images += (
                    '<div class="plot-container">'
                    f'<img src="{kbar}" alt="{label} K-value bar chart" '
                    'style="max-width: 100%; height: auto;"></div>'
                )
            blocks += f'<div style="page-break-before: auto;"><h3>{label}</h3>{images}</div>'

        if not blocks:
            return ""
        return (
            '<div style="page-break-before: auto;">'
            '<h2>Individual Sample Plots</h2>'
            f'{blocks}</div>'
        )

    def _create_comparison_k_distribution(self, sample_details, comparison_snapshot,
                                          breakdown: Optional[str] = None,
                                          plot_style=None) -> str:
        """Lognormal K-distribution comparison (histogram), matching the Comparison tab.

        Pools the report's K-values across the comparison scope and renders the
        same lognormal histogram view the Comparison tab shows (frequency axis,
        N labels, collapsed empty bins by default). Returns "" when no K-values
        are available.
        """
        if comparison_snapshot is None or not comparison_snapshot.k.included_records:
            return ""
        spec = self._build_comparison_spec(
            sample_details, comparison_snapshot,
            plot_type="k-distribution", breakdown=breakdown, plot_style=plot_style,
        )
        return _get_plot_export().export_comparison_spec(
            spec,
            figsize=REPORT_LANDSCAPE_FIGSIZE if spec.dense_report_layout else (12, 7),
        )

    def _create_k_value_boxplot(self, k_results_dict: Dict[str, List[KCalculationResult]]) -> str:
        """Create box plots for K-value comparison across samples."""
        if not k_results_dict:
            return ""
        return _get_plot_export().export_k_boxplot(k_results_dict)

    def _comparison_uses_grouped_k_scope(self, comparison_snapshot) -> bool:
        return any(group != UNGROUPED_LABEL for group in comparison_snapshot.k.group_names)

    def _k_scope_plot_colors(self, comparison_snapshot, series) -> List[str]:
        return k_scope_plot_colors(comparison_snapshot.k.group_names, series)

    def _create_comparison_k_scope_boxplot(self, comparison_snapshot, plot_style=None) -> str:
        """Create the report K boxplot from the shared comparison aggregation."""
        series = k_scope_value_series(comparison_snapshot.k)
        if not any(values for _label, values in series):
            return ""

        grouped = self._comparison_uses_grouped_k_scope(comparison_snapshot)
        title = (
            "Hydraulic Conductivity Distribution by Group"
            if grouped
            else "Hydraulic Conductivity Distribution by Dataset"
        )
        sample_count = comparison_snapshot.dataset_count
        style, _automatic_legend = self._comparison_report_style(
            plot_style, sample_count
        )
        kwargs = {"style": style}
        if sample_count >= REPORT_LARGE_BATCH_MIN_SAMPLES:
            kwargs["figsize"] = REPORT_LANDSCAPE_FIGSIZE
        return _get_plot_export().export_k_scope_boxplot(
            series,
            colors=self._k_scope_plot_colors(comparison_snapshot, series),
            title=title,
            **kwargs,
        )

    def _create_method_reliability_matrix(self, k_results_dict: Dict[str, List[KCalculationResult]],
                                          plot_style=None) -> str:
        """Create method reliability matrix for comparison report."""
        if not k_results_dict:
            return ""
        sample_count = len(k_results_dict)
        style, _automatic_legend = self._comparison_report_style(
            plot_style, sample_count
        )
        kwargs = {"style": style}
        if sample_count >= REPORT_LARGE_BATCH_MIN_SAMPLES:
            kwargs["figsize"] = REPORT_LANDSCAPE_FIGSIZE
        return _get_plot_export().export_reliability_matrix(k_results_dict, **kwargs)

    def _format_metadata_section(self, metadata: Dict[str, str]) -> str:
        """Format project metadata section with modern grid layout"""
        html = '<div class="metadata"><div class="metadata-grid">'

        if metadata.get('project_name'):
            html += '<div class="metadata-label">Project:</div>'
            html += f'<div class="metadata-value">{self._esc(metadata["project_name"])}</div>'
        if metadata.get('location'):
            html += '<div class="metadata-label">Location:</div>'
            html += f'<div class="metadata-value">{self._esc(metadata["location"])}</div>'
        if metadata.get('client'):
            html += '<div class="metadata-label">Client:</div>'
            html += f'<div class="metadata-value">{self._esc(metadata["client"])}</div>'
        if metadata.get('analyst'):
            html += '<div class="metadata-label">Analyst:</div>'
            html += f'<div class="metadata-value">{self._esc(metadata["analyst"])}</div>'

        html += '<div class="metadata-label">Report Date:</div>'
        html += f'<div class="metadata-value">{datetime.now().strftime("%B %d, %Y at %H:%M")}</div>'
        html += '</div></div>'

        return html

    def generate_grain_size_report(self, dataset: GrainSizeData,
                                  metadata: Optional[Dict[str, str]] = None,
                                  sections: Optional[Dict[str, bool]] = None,
                                  report_template: str = "standard",
                                  brand=None,
                                  appendix_label_config: Optional[Any] = None,
                                  plot_context: Optional[Dict[str, Any]] = None,
                                  k_results: Optional[List[KCalculationResult]] = None,
                                  selected_plots: Optional[set] = None) -> str:
        """
        Generate a grain size analysis report for a single sample

        Args:
            dataset: GrainSizeData object containing the sample data
            metadata: Dictionary with project metadata (project_name, location, client, analyst, notes)
            sections: Dictionary controlling which sections to include
            report_template: Template style - "standard", "executive", "technical", "appendix"

        Returns:
            HTML string of the complete report
        """

        # Set defaults
        if metadata is None:
            metadata = {}
        if sections is None:
            # Default sections based on template
            if report_template == "executive":
                sections = {
                    'cover_page': True,
                    'executive_summary': True,
                    'methodology': False,
                    'results': True,
                    'plots': True,
                    'raw_data': False,
                    'interpretation': True,
                    'percentiles': False,
                    'gradation': True,
                    'data_quality': False
                }
            elif report_template == "technical":
                sections = {
                    'cover_page': True,
                    'executive_summary': True,
                    'methodology': True,
                    'results': True,
                    'plots': True,
                    'raw_data': False,
                    'interpretation': True,
                    'percentiles': True,
                    'gradation': True,
                    'data_quality': True
                }
            elif report_template == "appendix":
                sections = {
                    'cover_page': False,
                    'executive_summary': False,
                    'methodology': False,
                    'results': False,
                    'plots': True,
                    'raw_data': True,
                    'interpretation': False,
                    'percentiles': True,
                    'gradation': True,
                    'data_quality': True
                }
            else:  # standard
                sections = {
                    'cover_page': False,
                    'executive_summary': True,
                    'methodology': True,
                    'results': True,
                    'plots': True,
                    'raw_data': False,
                    'interpretation': True,
                    'percentiles': True,
                    'gradation': True,
                    'data_quality': False
                }

        # Get characteristic grain sizes
        d10 = dataset.get_d10()
        d20 = dataset.get_d20()
        d30 = dataset.get_d30()
        d50 = dataset.get_d50()
        d60 = dataset.get_d60()

        # Calculate coefficients
        cu = (d60 / d10) if (d10 and d60 and d10 > 0) else None
        cc = ((d30 * d30) / (d10 * d60)) if (d10 and d30 and d60 and d10 > 0 and d60 > 0) else None

        # Start HTML report
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Grain Size Analysis Report - {dataset.sample_name}</title>
    {self._get_branded_style(brand)}
</head>
<body>
<div class="report-top-bar"></div>
"""

        # Cover Page (optional)
        if sections.get('cover_page', False):
            html += self._create_cover_page(
                "Grain Size Analysis",
                f"Sample: {dataset.sample_name}",
                metadata, brand
            )

        # Main Title (if no cover page)
        if not sections.get('cover_page', False):
            html += f'<h1>Grain Size Analysis Report</h1>'
            html += self._format_metadata_section(metadata)

        # Sample Information
        html += f"""
<div class="metadata">
    <div class="metadata-grid">
        <div class="metadata-label">Sample Name:</div>
        <div class="metadata-value">{dataset.sample_name}</div>
        <div class="metadata-label">Soil Classification:</div>
        <div class="metadata-value"><strong>{dataset.classify(scheme=self._scheme).label}</strong></div>
        <div class="metadata-label">Temperature:</div>
        <div class="metadata-value">{dataset.temperature}°C</div>
        <div class="metadata-label">Porosity:</div>
        <div class="metadata-value">{self._porosity_display(dataset)}</div>
        <div class="metadata-label">Porosity Source:</div>
        <div class="metadata-value">{self._esc(self._porosity_source_display(dataset))}</div>
        <div class="metadata-label">Data Points:</div>
        <div class="metadata-value">{len(dataset.particle_sizes)}</div>
    </div>
</div>
"""

        # Executive Summary
        if sections.get('executive_summary', True):
            html += f"""
<div class="no-break">
<h2>Executive Summary</h2>
<div class="success-box">
    <p>Sample <strong>{dataset.sample_name}</strong> has been analyzed and classified as <strong>{dataset.classify(scheme=self._scheme).label}</strong>.</p>
</div>
<div class="summary-stats">
    <div class="stat-card">
        <div class="stat-label">Median Size (D₅₀)</div>
        <div class="stat-value">{f'{d50:.3f}' if d50 else 'N/A'}</div>
        <div class="stat-unit">mm</div>
    </div>
    <div class="stat-card">
        <div class="stat-label">Uniformity (Cu)</div>
        <div class="stat-value">{f'{cu:.2f}' if cu else 'N/A'}</div>
        <div class="stat-unit">{self._classify_uniformity(cu)}</div>
    </div>
    <div class="stat-card">
        <div class="stat-label">Curvature (Cc)</div>
        <div class="stat-value">{f'{cc:.2f}' if cc else 'N/A'}</div>
        <div class="stat-unit">{self._classify_curvature(cc)}</div>
    </div>
</div>
</div>
"""

        # Methodology
        if sections.get('methodology', True):
            html += """
<div class="page-break">
<h2>Methodology</h2>
<div class="info-box">
    <h3>Grain Size Distribution Analysis</h3>
    <p>Grain size distribution was determined using sieve analysis and/or sedimentation methods following standard geotechnical procedures. The cumulative distribution curve plots percent passing versus grain size on a semi-logarithmic scale.</p>
</div>
<div class="info-box">
    <h3>Characteristic Diameters</h3>
    <p>Characteristic grain sizes represent specific percentiles of the grain size distribution:</p>
    <ul>
        <li><strong>D₁₀, D₃₀, D₅₀, D₆₀:</strong> Grain diameters at which 10%, 30%, 50%, and 60% of the soil mass is finer</li>
        <li><strong>D₅₀ (Median):</strong> The median grain size, representing the center of the distribution</li>
        <li>These values are fundamental for soil classification and hydraulic conductivity estimation</li>
    </ul>
</div>
<div class="info-box">
    <h3>Gradation Parameters</h3>
    <p><strong>Uniformity Coefficient:</strong> <code>Cu = D₆₀ / D₁₀</code></p>
    <ul>
        <li>Cu &lt; 4: Uniform gradation (narrow size range)</li>
        <li>4 ≤ Cu &lt; 6: Moderate gradation</li>
        <li>Cu ≥ 6: Well-graded soil (wide size range)</li>
    </ul>
    <p><strong>Coefficient of Curvature:</strong> <code>Cc = (D₃₀)² / (D₁₀ × D₆₀)</code></p>
    <ul>
        <li>1 ≤ Cc ≤ 3: Well-graded with good particle size distribution</li>
        <li>Outside this range: Gap-graded or uniform distribution</li>
    </ul>
</div>
</div>
"""

        # Results
        if sections.get('results', True):
            html += f"""
<div class="page-break">
<h2>Results & Analysis</h2>

<h3>Characteristic Grain Sizes</h3>
<div class="summary-stats">
    <div class="stat-card">
        <div class="stat-label">D₁₀</div>
        <div class="stat-value">{f'{d10:.3f}' if d10 else 'N/A'}</div>
        <div class="stat-unit">mm</div>
    </div>
    <div class="stat-card">
        <div class="stat-label">D₃₀</div>
        <div class="stat-value">{f'{d30:.3f}' if d30 else 'N/A'}</div>
        <div class="stat-unit">mm</div>
    </div>
    <div class="stat-card">
        <div class="stat-label">D₅₀</div>
        <div class="stat-value">{f'{d50:.3f}' if d50 else 'N/A'}</div>
        <div class="stat-unit">mm</div>
    </div>
    <div class="stat-card">
        <div class="stat-label">D₆₀</div>
        <div class="stat-value">{f'{d60:.3f}' if d60 else 'N/A'}</div>
        <div class="stat-unit">mm</div>
    </div>
</div>

<h3>Soil Classification Parameters</h3>
<table data-report-table="soil-classification">
    <thead>
        <tr>
            <th>Parameter</th>
            <th>Value</th>
            <th>Classification</th>
        </tr>
    </thead>
    <tbody>
        <tr>
            <td><strong>Uniformity Coefficient</strong> (Cu = D₆₀/D₁₀)</td>
            <td class="text-center">{f'{cu:.2f}' if cu else 'N/A'}</td>
            <td><span class="badge badge-info">{self._classify_uniformity(cu)}</span></td>
        </tr>
        <tr>
            <td><strong>Coefficient of Curvature</strong> (Cc = D₃₀²/D₁₀·D₆₀)</td>
            <td class="text-center">{f'{cc:.2f}' if cc else 'N/A'}</td>
            <td><span class="badge badge-info">{self._classify_curvature(cc)}</span></td>
        </tr>
        <tr>
            <td><strong>Soil Classification</strong></td>
            <td colspan="2"><span class="badge badge-success">{dataset.classify(scheme=self._scheme).label}</span></td>
        </tr>
    </tbody>
</table>
"""

            if sections.get('gradation', True):
                html += f"<h3>Gradation Analysis</h3>{self._create_gradation_table(dataset)}"

            html += "</div>"

        # Visual Charts — each plot is individually selectable.
        if sections.get('plots', True):
            selected = (
                selected_plots if selected_plots is not None
                else {'grain_size_curve', 'k_value_bar'}
            )
            figure_no = 1

            if 'grain_size_curve' in selected:
                grain_plot = self._create_grain_size_plot(dataset, plot_context)
                html += f"""
<div class="page-break">
<h2>Grain Size Distribution Curve</h2>
<div class="plot-container">
    <img src="{grain_plot}" alt="Grain Size Distribution" />
    <div class="figure-caption">Figure {figure_no}: Cumulative grain size distribution curve for {dataset.sample_name}</div>
</div>
</div>
"""
                figure_no += 1

            if 'grain_size_histogram' in selected:
                histogram_plot = self._create_grain_size_histogram(dataset, plot_context)
                html += f"""
<div class="page-break">
<h2>Grain-size Class Histogram</h2>
<div class="plot-container">
    <img src="{histogram_plot}" alt="Grain-size class histogram" />
    <div class="figure-caption">Figure {figure_no}: Retained weight by classification class for {dataset.sample_name}</div>
</div>
</div>
"""
                figure_no += 1

            if k_results and 'k_value_bar' in selected:
                k_bar_chart = self._create_k_value_bar_chart(k_results, plot_context)
                if k_bar_chart:
                    html += f"""
<div class="page-break">
<h2>Hydraulic Conductivity by Method</h2>
<div class="plot-container">
    <img src="{k_bar_chart}" alt="K-Value Bar Chart" />
    <div class="figure-caption">Figure {figure_no}: Estimated hydraulic conductivity by method for {dataset.sample_name}</div>
</div>
</div>
"""
                    figure_no += 1

            if k_results and 'applicability_heatmap' in selected:
                method_heatmap = self._create_method_applicability_heatmap(k_results, plot_context)
                if method_heatmap:
                    html += f"""
<div class="page-break">
<h2>Method Applicability Status</h2>
<div class="plot-container">
    <img src="{method_heatmap}" alt="Method Applicability Heatmap" />
    <div class="figure-caption">Figure {figure_no}: Method applicability for {dataset.sample_name}</div>
</div>
</div>
"""
                    figure_no += 1

        # Interpretation
        if sections.get('interpretation', True):
            html += f"""
<div class="page-break">
<h2>Interpretation & Discussion</h2>
<div class="info-box">
    <h3>Grain Size Distribution Analysis</h3>
    <p>{self._interpret_grain_distribution(dataset, cu, cc)}</p>
</div>
"""
            if metadata.get('notes'):
                html += f"""
<div class="info-box">
    <h3>Additional Notes</h3>
    <p>{self._note_html(metadata['notes'])}</p>
</div>
"""
            html += "</div>"

        appendix_document = self._build_grain_size_appendix_document(
            dataset,
            sections,
            appendix_label_config=appendix_label_config,
        )
        html += self._render_appendix_document_html(appendix_document)

        # Footer
        html += """
<div class="footer">
    <span><strong>Grain Size Analysis Report</strong></span>
    <span>Generated by Grain Size Analysis &amp; Hydraulic Conductivity Calculator</span>
</div>
</body>
</html>
"""

        return html

    def generate_k_value_report(self, dataset: GrainSizeData,
                               k_results: List[KCalculationResult],
                               temperature: float,
                               porosity: float,
                               metadata: Optional[Dict[str, str]] = None,
                               sections: Optional[Dict[str, bool]] = None,
                               report_template: str = "standard",
                               brand=None) -> str:
        """
        Generate a hydraulic conductivity (K-value) report for a single sample

        Args:
            dataset: GrainSizeData object
            k_results: List of KCalculationResult objects from different methods
            temperature: Water temperature in °C
            porosity: Soil porosity
            metadata: Project metadata dictionary
            sections: Dictionary controlling which sections to include
            report_template: Template style - "standard", "executive", "technical", "appendix"

        Returns:
            HTML string of the complete report
        """

        # Set defaults
        if metadata is None:
            metadata = {}
        if sections is None:
            if report_template == "executive":
                sections = {
                    'cover_page': True,
                    'executive_summary': True,
                    'methodology': False,
                    'results': True,
                    'plots': True,
                    'interpretation': True,
                    'k_statistics': False
                }
            elif report_template == "technical":
                sections = {
                    'cover_page': True,
                    'executive_summary': True,
                    'methodology': True,
                    'results': True,
                    'plots': True,
                    'interpretation': True,
                    'k_statistics': True
                }
            elif report_template == "appendix":
                sections = {
                    'cover_page': False,
                    'executive_summary': False,
                    'methodology': False,
                    'results': False,
                    'plots': True,
                    'interpretation': False,
                    'k_statistics': True
                }
            else:  # standard
                sections = {
                    'cover_page': False,
                    'executive_summary': True,
                    'methodology': True,
                    'results': True,
                    'plots': True,
                    'interpretation': True,
                    'k_statistics': True
                }

        summary = build_k_result_summary(k_results)

        if summary.geometric_mean_m_s is None:
            return self._generate_no_results_report(dataset.sample_name)

        # Shared OK-only K summary.
        mean_k = summary.geometric_mean_m_s
        arithmetic_k = summary.arithmetic_mean_m_s
        median_k = summary.median_m_s
        std_k = summary.std_dev_m_s or 0.0
        min_k = summary.min_m_s
        max_k = summary.max_m_s
        variability_ratio = max_k / min_k if min_k > 0 else 0

        # Start HTML report
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Hydraulic Conductivity Report - {dataset.sample_name}</title>
    {self._get_branded_style(brand)}
</head>
<body>
<div class="report-top-bar"></div>
"""

        # Cover Page (optional)
        if sections.get('cover_page', False):
            html += self._create_cover_page(
                "Hydraulic Conductivity Analysis",
                f"Sample: {dataset.sample_name}",
                metadata, brand
            )

        # Main Title (if no cover page)
        if not sections.get('cover_page', False):
            html += f'<h1>Hydraulic Conductivity Analysis Report</h1>'
            html += self._format_metadata_section(metadata)

        # Sample Information
        html += f"""
<div class="metadata">
    <div class="metadata-grid">
        <div class="metadata-label">Sample Name:</div>
        <div class="metadata-value">{dataset.sample_name}</div>
        <div class="metadata-label">Temperature:</div>
        <div class="metadata-value">{temperature}°C</div>
        <div class="metadata-label">Porosity:</div>
        <div class="metadata-value">{self._porosity_display(dataset, porosity)}</div>
        <div class="metadata-label">Porosity Source:</div>
        <div class="metadata-value">{self._esc(self._porosity_source_display(dataset))}</div>
        <div class="metadata-label">Methods Evaluated:</div>
        <div class="metadata-value">{len(k_results)} empirical methods</div>
        <div class="metadata-label">Valid Results:</div>
        <div class="metadata-value"><span class="badge badge-success">{summary.included_count} / {summary.total_cells}</span></div>
    </div>
</div>
"""

        # Executive Summary
        if sections.get('executive_summary', True):
            html += f"""
            <div style="page-break-before: auto;">
            <h2>Executive Summary</h2>
            <div class="info-box">
                <p><strong>Sample:</strong> {dataset.sample_name} hydraulic conductivity analysis using {len(k_results)} empirical methods.</p>
                <p><strong>Geometric Mean K:</strong> {mean_k:.2e} m/s (from {summary.included_count} OK methods)</p>
                <p><strong>Permeability Classification:</strong> {self._classify_permeability(mean_k)}</p>
                <p><strong>Variability:</strong> {max_k/min_k:.1f}x difference between minimum and maximum estimates</p>
            </div>
            </div>
            """

        # Methodology
        if sections.get('methodology', True):
            html += """
            <div style="page-break-before: auto;">
            <h2>Methodology</h2>
            <div class="info-box">
                <h3>Hydraulic Conductivity Estimation</h3>
                <p>Hydraulic conductivity (K) represents the ease with which water can move through pore spaces
                in soil. This analysis employs multiple empirical methods developed from various grain size
                parameters to estimate K-values for comparison and reliability assessment.</p>
                <h3>Empirical Methods</h3>
                <p>Each method has specific applicability ranges and underlying assumptions based on soil type,
                grain size distribution, and original calibration data. Methods include Hazen, Shepherd, Kozeny-Carman,
                Terzaghi, Breyer, Slichter, Sauerbrei, Kruger, Zunker, Zamarin, USBR, and Barr.</p>
                <h3>Quality Assessment</h3>
                <p>Each calculation is evaluated for applicability based on grain size parameters. Results are
                marked as OK (within recommended range), WARNING (outside optimal range), or ERROR (calculation failed).
                Statistical analysis of multiple methods provides confidence bounds on the estimated K-value.</p>
            </div>
            </div>
            """

        # Results
        if sections.get('results', True):
            html += f"""
            <div style="page-break-before: auto;">
            <h2>Results & Analysis</h2>

            <h3>Statistical Summary</h3>
            <div class="summary-stats">
                <div class="stat-card">
                    <div class="stat-label">K Geometric Mean</div>
                    <div class="stat-value">{mean_k:.2e} m/s</div>
                </div>
                <div class="stat-card">
                    <div class="stat-label">K Arithmetic Mean</div>
                    <div class="stat-value">{arithmetic_k:.2e} m/s</div>
                </div>
                <div class="stat-card">
                    <div class="stat-label">Median K</div>
                    <div class="stat-value">{median_k:.2e} m/s</div>
                </div>
                <div class="stat-card">
                    <div class="stat-label">Min K</div>
                    <div class="stat-value">{min_k:.2e} m/s</div>
                </div>
                <div class="stat-card">
                    <div class="stat-label">Max K</div>
                    <div class="stat-value">{max_k:.2e} m/s</div>
                </div>
            </div>

            <h3>K-Value Calculations by Method</h3>
            <table class="table-compact" data-report-table="k-method-results">
                <tr>
                    <th>Method</th>
                    <th class="num">K-Value (m/s)</th>
                    <th>Formula</th>
                    <th>Status</th>
                </tr>
            """

            for result in k_results:
                k_display = f"{result.k_value:.2e}" if result.k_value else "N/A"

                html += f"""
                <tr>
                    <td>{result.method_name}</td>
                    <td class="num">{k_display}</td>
                    <td class="note">{self._esc(result.formula_used)}</td>
                    <td>{self._esc(result.status_message or result.status)}</td>
                </tr>
                """

            html += """
            </table>

            <h3>Permeability Classification</h3>
            <div class="info-box">
                <p><strong>Classification:</strong> {}</p>
                <p><strong>Typical Application:</strong> {}</p>
            </div>
            """.format(self._classify_permeability(mean_k), self._get_permeability_application(mean_k))

            html += "</div>"

        elif sections.get('k_statistics', True):
            html += f"""
            <div style="page-break-before: auto;">
            <h2>K-Value Results</h2>
            <h3>K-Value Calculations by Method</h3>
            {self._create_k_statistics_table(k_results)}
            </div>
            """

        # Visual Charts
        if sections.get('plots', True):
            k_bar_chart = self._create_k_value_bar_chart(k_results, plot_context)
            method_heatmap = self._create_method_applicability_heatmap(k_results, plot_context)

            if k_bar_chart:
                html += f"""
                <div style="page-break-before: auto;">
                <h2>K-Value Comparison Chart</h2>
                <div class="plot-container">
                    <img src="{k_bar_chart}" alt="K-Value Bar Chart" style="max-width: 100%; height: auto;">
                </div>
                </div>
                """

            if method_heatmap:
                html += f"""
                <div style="page-break-before: auto;">
                <h2>Method Applicability Status</h2>
                <div class="plot-container">
                    <img src="{method_heatmap}" alt="Method Applicability Heatmap" style="max-width: 100%; height: auto;">
                </div>
                </div>
                """

        # Interpretation
        if sections.get('interpretation', True):
            html += f"""
            <div style="page-break-before: auto;">
            <h2>Interpretation & Discussion</h2>
            <div class="info-box">
                <h3>Method Variability Analysis</h3>
                <p><strong>Variability:</strong> {max_k/min_k:.1f}x difference between min and max</p>
                <p><strong>Standard Deviation:</strong> {std_k:.2e} m/s</p>
                <p><strong>Coefficient of Variation:</strong> {(std_k/mean_k)*100:.1f}%</p>
                <p>{self._interpret_k_variability(max_k/min_k)}</p>
            </div>
            """

            # Add custom notes if provided
            if metadata.get('notes'):
                html += f"""
                <div class="info-box">
                    <h3>Additional Notes</h3>
                    <p>{self._note_html(metadata['notes'])}</p>
                </div>
                """

            html += "</div>"

        # Add footer
        html += """
            <div class="footer">
                <span><strong>Hydraulic Conductivity Report</strong></span>
                <span>Generated by Grain Size Analysis &amp; Hydraulic Conductivity Calculator</span>
            </div>
        </body>
        </html>
        """

        return html

    def generate_combined_report(self, dataset: GrainSizeData,
                                k_results: List[KCalculationResult],
                                temperature: float,
                                porosity: float,
                                metadata: Optional[Dict[str, str]] = None,
                                sections: Optional[Dict[str, bool]] = None,
                                brand=None,
                                appendix_label_config: Optional[Any] = None,
                                plot_context: Optional[Dict[str, Any]] = None) -> str:
        """Generate a combined report with both grain size and K-value analysis"""

        # Set defaults
        if metadata is None:
            metadata = {}
        if sections is None:
            sections = {
                'executive_summary': True,
                'methodology': True,
                'results': True,
                'plots': True,
                'raw_data': False,
                'interpretation': True
            }

        grain_sections = dict(sections)
        k_sections = dict(sections)
        k_sections['cover_page'] = False

        grain_report = self.generate_grain_size_report(
            dataset,
            metadata=metadata,
            sections=grain_sections,
            brand=brand,
            appendix_label_config=appendix_label_config,
            plot_context=plot_context,
        )
        k_report = self.generate_k_value_report(
            dataset,
            k_results,
            temperature,
            porosity,
            metadata=metadata,
            sections=k_sections,
            brand=brand,
        )

        grain_body = self._extract_body_contents(grain_report)
        k_body = self._strip_leading_report_markup(self._extract_body_contents(k_report))

        # Create combined report with page break between sections
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Complete Analysis Report - {dataset.sample_name}</title>
            {self._get_branded_style(brand)}
        </head>
        <body>
            {grain_body.replace('</body>', '').replace('</html>', '')}

            <div class="page-break"></div>

            <h1>Hydraulic Conductivity Analysis</h1>
            {k_body}
        </body>
        </html>
        """

        return html

    def generate_comparison_report(self, datasets: List[GrainSizeData],
                                  k_results_dict: Optional[Dict[str, List[KCalculationResult]]] = None,
                                  temperature: Optional[float] = None,
                                  porosity: Optional[float] = None,
                                  metadata: Optional[Dict[str, str]] = None,
                                  sections: Optional[Dict[str, bool]] = None,
                                  brand=None,
                                  sample_details: Optional[List[Dict[str, Any]]] = None,
                                  selected_plots: Optional[set] = None,
                                  plot_breakdowns: Optional[Dict[str, str]] = None,
                                  plot_style=None,
                                  progress=None,
                                  cancel_check=None) -> str:
        """Generate a comparison report for multiple samples.

        *plot_breakdowns* optionally maps a comparison plot key
        (``distribution_overlay`` / ``k_value_comparison``) to ``"group"`` or
        ``"dataset"``; omitted keys fall back to auto (group when groups exist).
        *plot_style* is the global report/export ``PlotStyle`` themed once on the
        report tab; ``None`` keeps the default preset.

        *progress* (optional ``callable(current, total, label)``) and
        *cancel_check* (optional ``callable() -> bool``) let a worker thread show
        progress and abort: the per-sample plot loop reports per sample and raises
        :class:`ReportCancelled` when cancellation is requested.
        """
        if metadata is None:
            metadata = {}
        if sections is None:
            sections = {
                'cover_page': True,
                'executive_summary': True,
                'methodology': True,
                'results': True,
                'plots': True,
                'interpretation': True,
                'grain_comparison': True,
                'k_statistics': True,
            }

        if sample_details is None:
            k_results_dict = k_results_dict or {}
            sample_details = []
            for index, dataset in enumerate(datasets):
                label = dataset.sample_name or f"Sample {index + 1}"
                sample_details.append({
                    "label": label,
                    "dataset": dataset,
                    "k_results": list(k_results_dict.get(label, [])),
                    "group_name": getattr(dataset, "group_name", "Ungrouped"),
                    "temperature": temperature,
                    "porosity": porosity,
                    "plot_context": None,
                })
        else:
            normalized_details = []
            for index, item in enumerate(sample_details):
                dataset = item.get("dataset")
                if dataset is None:
                    continue
                normalized_details.append({
                    "label": item.get("label") or dataset.sample_name or f"Sample {index + 1}",
                    "dataset": dataset,
                    "k_results": list(item.get("k_results") or []),
                    "group_name": item.get("group_name") or getattr(dataset, "group_name", "Ungrouped"),
                    "temperature": item.get("temperature"),
                    "porosity": item.get("porosity"),
                    "plot_context": item.get("plot_context"),
                })
            sample_details = normalized_details

        datasets = [item["dataset"] for item in sample_details]
        sample_labels = [str(item["label"]) for item in sample_details]
        plot_results_dict = {
            str(item["label"]): list(item.get("k_results") or [])
            for item in sample_details
        }
        snapshot_inputs = [
            DatasetAnalysisInput(
                label=str(item["label"]),
                dataset=item["dataset"],
                k_results=tuple(item.get("k_results") or ()),
                group_name=item.get("group_name") or getattr(item["dataset"], "group_name", "Ungrouped"),
                temperature=item.get("temperature"),
                porosity=item.get("porosity"),
                plot_context=item.get("plot_context"),
            )
            for item in sample_details
        ]
        comparison_snapshot = build_comparison_snapshot(
            snapshot_inputs,
            ComparisonSnapshotOptions(
                k_options=KAggregationOptions(include_warnings=False),
                classification_scheme=self._scheme,
            ),
        )
        mean_k_by_sample = {
            name: stats.geometric_mean_m_s
            for name, stats in comparison_snapshot.k.by_dataset.items()
            if stats.geometric_mean_m_s is not None
        }

        temperature_summary = self._summarize_sample_field(sample_details, "temperature", " °C")
        porosity_summary = self._summarize_sample_field(sample_details, "porosity")

        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Multi-Sample Comparison Report</title>
            {self._get_branded_style(brand)}
        </head>
        <body>
        <div class="report-top-bar"></div>
        """

        if sections.get('cover_page', True):
            html += self._create_cover_page(
                title="Multi-Sample Comparison Report",
                subtitle=f"Hydraulic Conductivity Analysis - {len(datasets)} Samples",
                metadata=metadata,
                brand=brand,
            )

        html += f"""
            <h1>Multi-Sample Comparison Report</h1>

            {self._format_metadata_section(metadata)}

            <div class="metadata">
                <p><strong>Number of Samples:</strong> {len(datasets)}</p>
                <p><strong>Temperature:</strong> {self._esc(temperature_summary)}</p>
                <p><strong>Porosity:</strong> {self._esc(porosity_summary)}</p>
            </div>
        """

        if sections.get('executive_summary', True):
            html += f"""
            <div style="page-break-before: auto;">
            <h2>Executive Summary</h2>
            <div class="info-box">
                <p><strong>Comparison Analysis:</strong> This report compares {len(datasets)} soil samples
                based on grain size distribution and hydraulic conductivity estimates.</p>
            """

            if mean_k_by_sample:
                highest = max(mean_k_by_sample.items(), key=lambda x: x[1])
                lowest = min(mean_k_by_sample.items(), key=lambda x: x[1])
                html += f"""
                <p><strong>Key Findings:</strong></p>
                <ul>
                    <li>Highest permeability: {self._esc(highest[0])} ({highest[1]:.2e} m/s)</li>
                    <li>Lowest permeability: {self._esc(lowest[0])} ({lowest[1]:.2e} m/s)</li>
                    <li>Permeability range: {highest[1]/lowest[1]:.1f}x difference</li>
                </ul>
                """

            html += """
            </div>
            </div>
            """

        if sections.get('methodology', True):
            html += """
            <div style="page-break-before: auto;">
            <h2>Methodology</h2>
            <div class="info-box">
                <h3>Comparative Analysis Approach</h3>
                <p>This comparison report presents a side-by-side analysis of multiple soil samples
                to identify patterns, variations, and relationships between grain size characteristics
                and hydraulic conductivity estimates.</p>
                <h3>Analysis Components</h3>
                <p><strong>Grain Size Comparison:</strong> Overlapping distribution curves allow visual
                assessment of particle size variations between samples.</p>
                <p><strong>K-Value Comparison:</strong> Box plots and statistical summaries reveal the
                range and reliability of hydraulic conductivity estimates across samples.</p>
                <p><strong>Method Reliability:</strong> A reliability matrix shows which empirical methods
                are applicable for each sample, helping identify the most suitable estimation approaches.</p>
            </div>
            </div>
            """

        if sections.get('results', True):
            html += """
            <div style="page-break-before: auto;">
            <h2>Results & Analysis</h2>
            """
            html += (
                f"<h3>Sample Overview</h3>"
                f"{self._create_sample_overview_tables(sample_details, mean_k_by_sample)}"
            )

            if sections.get('grain_comparison', True):
                html += f"<h3>Grain Parameters Comparison</h3>{self._create_grain_parameters_comparison_table(datasets, sample_labels)}"

            if sections.get('k_statistics', True) and plot_results_dict:
                html += f"<h3>K-Value Aggregate Summary</h3>{self._create_comparison_k_scope_summary_table(comparison_snapshot)}"
                html += f"<h3>K-Value Calculations by Dataset and Method</h3>{self._create_comparison_k_statistics_table(plot_results_dict)}"
                html += f"<h3>Permeability Classification Summary</h3>{self._create_permeability_classification_table(plot_results_dict)}"

            html += "</div>"

        elif sections.get('k_statistics', True) and plot_results_dict:
            html += f"""
            <div style="page-break-before: auto;">
            <h2>K-Value Results</h2>
            <h3>K-Value Aggregate Summary</h3>
            {self._create_comparison_k_scope_summary_table(comparison_snapshot)}
            <h3>K-Value Calculations by Dataset and Method</h3>
            {self._create_comparison_k_statistics_table(plot_results_dict)}
            <h3>Permeability Classification Summary</h3>
            {self._create_permeability_classification_table(plot_results_dict)}
            </div>
            """

        if sections.get('plots', True):
            # Each comparison plot is individually selectable; the default keeps
            # the overlay, K-value bars and K boxplot and leaves the lognormal
            # K-distribution and reliability matrix opt-in.
            selected = (
                selected_plots if selected_plots is not None
                else {'distribution_overlay', 'k_value_comparison',
                      'statistical_boxplots', 'reliability_matrix'}
            )
            breakdowns = plot_breakdowns or {}

            # Progress / cancellation: the cross-sample plots are one coarse step;
            # the per-sample loop reports per sample (the part that scales with
            # dataset count).
            per_sample_on = bool({'per_sample_grain', 'per_sample_histogram', 'per_sample_kbar'} & selected)
            total_steps = 1 + (len(sample_details) if per_sample_on else 0) + 1
            step_state = {'n': 0}

            def _advance(label: str) -> None:
                if cancel_check is not None and cancel_check():
                    raise ReportCancelled()
                step_state['n'] += 1
                if progress is not None:
                    progress(min(step_state['n'], total_steps), total_steps, label)

            _advance("Rendering comparison plots")

            def _plot_variants(create_fn, key, note_fn=None):
                """Render a breakdown-capable plot, expanding "both" to two images.

                Returns ``(caption_suffix, data_uri, note)`` entries. A "both"
                breakdown emits the per-group and per-dataset variants as two
                entries; any other value yields a single entry.
                """
                if key not in selected:
                    return []
                chosen = breakdowns.get(key)
                if chosen == 'both':
                    wanted = [(' — per group', 'group'),
                              (' — per dataset', 'dataset')]
                else:
                    wanted = [('', chosen)]
                variants = []
                for suffix, bd in wanted:
                    uri = create_fn(bd)
                    if uri:
                        note = note_fn(bd) if note_fn is not None else ""
                        variants.append((suffix, uri, note))
                return variants

            large_batch_layout = len(sample_details) >= REPORT_LARGE_BATCH_MIN_SAMPLES

            def _plot_block(title, alt, variants, page_break='auto'):
                block = ""
                for variant in variants:
                    suffix, uri = variant[:2]
                    note = variant[2] if len(variant) > 2 else ""
                    note_html = (
                        f'<p class="figure-caption">{self._esc(note)}</p>'
                        if note
                        else ""
                    )
                    classes = "comparison-plot-page"
                    effective_break = page_break
                    if large_batch_layout:
                        classes += " landscape-plot-page"
                        effective_break = "always"
                    block += f"""
                <div class="{classes}" style="page-break-before: {effective_break};">
                <h2>{title}{suffix}</h2>
                {note_html}
                <div class="plot-container">
                    <img src="{uri}" alt="{alt}" style="max-width: 100%; height: auto;">
                </div>
                </div>
                """
                return block

            html += _plot_block(
                "Grain Size Distribution Comparison", "Grain Size Comparison",
                _plot_variants(
                    lambda bd: self._create_comparison_grain_size_plot(
                        sample_details, comparison_snapshot,
                        breakdown=bd, plot_style=plot_style),
                    'distribution_overlay'),
            )
            html += _plot_block(
                "Grain-size Class Fractions", "Grain-size Class Fractions",
                _plot_variants(
                    lambda bd: self._create_comparison_grain_size_histogram(
                        sample_details, comparison_snapshot,
                        breakdown=bd, plot_style=plot_style),
                    'grain_size_histogram_comparison',
                    note_fn=lambda bd: self._class_fraction_layout_note(
                        sample_details,
                        comparison_snapshot,
                        breakdown=bd,
                        plot_style=plot_style,
                    )),
            )
            html += _plot_block(
                "Hydraulic Conductivity by Method", "K-Value Comparison",
                _plot_variants(
                    lambda bd: self._create_comparison_k_value_bar(
                        sample_details, comparison_snapshot,
                        breakdown=bd, plot_style=plot_style),
                    'k_value_comparison'),
            )
            html += _plot_block(
                "Hydraulic Conductivity Distribution (Lognormal)", "K Distribution",
                _plot_variants(
                    lambda bd: self._create_comparison_k_distribution(
                        sample_details, comparison_snapshot,
                        breakdown=bd, plot_style=plot_style),
                    'k_distribution'),
            )

            if 'statistical_boxplots' in selected:
                k_boxplot = self._create_comparison_k_scope_boxplot(
                    comparison_snapshot, plot_style=plot_style)
                if k_boxplot:
                    html += _plot_block(
                        "Hydraulic Conductivity Distribution", "K-Value Boxplot",
                        [('', k_boxplot)])

            if 'reliability_matrix' in selected:
                reliability_matrix = self._create_method_reliability_matrix(
                    plot_results_dict, plot_style=plot_style)
                if reliability_matrix:
                    html += _plot_block(
                        "Appendix: Method Reliability Matrix", "Method Reliability Matrix",
                        [('', reliability_matrix)], page_break='always')

            # Per-sample (individual) plots — each sample's own grain curve and/or
            # K-value bar, so a multi-sample report can carry per-sample detail
            # alongside the cross-sample plots.
            # Resolve each sample's overlay colour (from the active palette) so its
            # per-sample grain curve matches the comparison overlay.
            per_sample_colors = {}
            if per_sample_on and {'per_sample_grain', 'per_sample_histogram'} & selected:
                color_spec = self._build_comparison_spec(
                    sample_details, comparison_snapshot,
                    plot_type="distribution", plot_style=plot_style,
                )
                per_sample_colors = dict(color_spec.color_by_name)
            html += self._create_per_sample_plots_section(
                sample_details,
                include_grain='per_sample_grain' in selected,
                include_histogram='per_sample_histogram' in selected,
                include_kbar='per_sample_kbar' in selected,
                advance=_advance if per_sample_on else None,
                colors_by_name=per_sample_colors,
            )
            _advance("Finalizing report")

        if sections.get('interpretation', True):
            html += """
            <div style="page-break-before: auto;">
            <h2>Interpretation & Discussion</h2>
            <div class="info-box">
                <h3>Comparative Analysis</h3>
            """

            if mean_k_by_sample:
                highest = max(mean_k_by_sample.items(), key=lambda x: x[1])
                lowest = min(mean_k_by_sample.items(), key=lambda x: x[1])

                html += f"""
                <p><strong>Permeability Characteristics:</strong></p>
                <ul>
                    <li>The highest permeability sample is {self._esc(highest[0])} with K = {highest[1]:.2e} m/s,
                    classified as {self._classify_permeability(highest[1])}.</li>
                    <li>The lowest permeability sample is {self._esc(lowest[0])} with K = {lowest[1]:.2e} m/s,
                    classified as {self._classify_permeability(lowest[1])}.</li>
                    <li>The {highest[1]/lowest[1]:.1f}-fold difference in permeability reflects the
                    variability in grain size distribution among the samples.</li>
                </ul>
                """

                all_k_values = list(mean_k_by_sample.values())
                mean_all = np.mean(all_k_values)
                std_all = np.std(all_k_values)

                html += f"""
                <p><strong>Statistical Overview:</strong></p>
                <ul>
                    <li>Average sample K geometric mean: {mean_all:.2e} m/s</li>
                    <li>Standard deviation: {std_all:.2e} m/s</li>
                    <li>Coefficient of variation: {(std_all/mean_all)*100:.1f}%</li>
                </ul>
                """

            html += """
            </div>
            """

            if metadata.get('notes'):
                html += f"""
                <div class="info-box">
                    <h3>Additional Notes</h3>
                    <p>{self._note_html(metadata['notes'])}</p>
                </div>
                """

            html += "</div>"

        html += """
            <div class="footer">
                <span><strong>Multi-Sample Comparison Report</strong></span>
                <span>Generated by Grain Size Analysis &amp; Hydraulic Conductivity Calculator</span>
            </div>
        </body>
        </html>
        """

        return html

    # Helper methods
    def _classify_uniformity(self, cu: Optional[float]) -> str:
        if cu is None:
            return "Cannot calculate"
        return _gc_cu_label(cu)

    def _classify_curvature(self, cc: Optional[float]) -> str:
        if cc is None:
            return "Cannot calculate"
        lbl = _gc_cc_label(cc)
        if lbl == "Well-graded range":
            return "Well-graded (1 \u2264 Cc \u2264 3)"
        return "Gap-graded or Uniform"

    def _create_percentiles_table(self, dataset: GrainSizeData) -> str:
        """Generate HTML table with percentiles (D5, D10, D16, D20, D25, D30, D40, D50, D60, D75, D84, D90, D95)"""
        percentiles_list = [5, 10, 16, 20, 25, 30, 40, 50, 60, 75, 84, 90, 95]
        percentiles_dict = self._calculate_percentile_values(dataset, percentiles_list)
        max_val = max((value for value in percentiles_dict.values() if value is not None), default=0)

        html = """
        <table data-report-table="percentiles">
            <thead>
                <tr>
                    <th>Percentile</th>
                    <th>Size (mm)</th>
                    <th>Relative (%)</th>
                </tr>
            </thead>
            <tbody>
        """

        for p in percentiles_list:
            val = percentiles_dict[p]
            bar_width = int((val / max_val) * 100) if val is not None and max_val > 0 else 0

            # Highlight key percentiles (D10, D30, D50, D60)
            is_key = p in [10, 30, 50, 60]

            html += f"""
            <tr>
                <td style="text-align: center;"><strong>D{p}{'*' if is_key else ''}</strong></td>
                <td style="text-align: right;">{f'{val:.3f}' if val is not None else 'N/A'}</td>
                <td style="text-align: right;">{bar_width}%</td>
            </tr>
            """

        html += "</tbody></table>"
        return html

    def _create_gradation_table(self, dataset: GrainSizeData, scheme=None) -> str:
        """Grain-size classification, detailed sub-class fractions, descriptor and
        calculation internals \u2014 respecting the active classification scheme."""
        s = scheme if scheme is not None else self._scheme

        try:
            result = dataset.classify(scheme=s)
        except Exception:
            return ""

        cu = (dataset.get_uniformity_coefficient()
              if hasattr(dataset, "get_uniformity_coefficient") else None)
        d50 = dataset.get_d50() if hasattr(dataset, "get_d50") else None
        descriptor = _gc_sed_descriptor(result.fractions, d50, cu, s)
        scheme_name = getattr(s, "name", "scheme")

        html = f"<p><strong>Label ({self._esc(scheme_name)}):</strong> {self._esc(result.label)}</p>"
        if descriptor:
            html += f"<p><strong>Descriptor:</strong> {self._esc(descriptor)}</p>"

        # Detailed, scheme-aware sub-class breakdown.
        html += """
        <table data-report-table="gradation">
            <thead>
                <tr><th>Sub-class</th><th>Size range (mm)</th><th>Percentage</th></tr>
            </thead>
            <tbody>
        """
        for d in (result.detailed_fractions or ()):
            rng = f"&lt; {d.upper_mm:g}" if d.lower_mm <= 0 else f"{d.lower_mm:g} &ndash; {d.upper_mm:g}"
            dominant = " class=\"row-emphasis\"" if d.label == result.detailed_class else ""
            html += (
                f"<tr{dominant}><td><strong>{self._esc(d.label)}</strong></td>"
                f"<td>{rng}</td>"
                f"<td style=\"text-align: right;\">{d.pct:.1f}%</td></tr>"
            )
        html += "</tbody></table>"

        html += self._create_calculation_internals_html(dataset)
        return html

    def _create_calculation_internals_html(self, dataset: GrainSizeData) -> str:
        """Render the intermediate K-calculation values (constants, effective
        diameters, phi/Folk-Ward, porosity functions) as small HTML tables."""
        porosity = (dataset.effective_porosity()
                    if hasattr(dataset, "effective_porosity")
                    else getattr(dataset, "porosity", None))
        internals = compute_calculation_internals(
            dataset.particle_sizes, dataset.percent_passing,
            getattr(dataset, "temperature", 20.0), porosity,
        )
        html = "<h4>Calculation Internals</h4>"
        for group in internals.groups():
            html += (
                f"<table data-report-table=\"calc-internals\">"
                f"<thead><tr><th>{self._esc(group.title)}</th><th>Value</th></tr></thead><tbody>"
            )
            for label, value in group.rows:
                html += (
                    f"<tr><td>{self._esc(label)}</td>"
                    f"<td style=\"text-align: right;\">{self._esc(value)}</td></tr>"
                )
            html += "</tbody></table>"
        return html

    def _create_k_statistics_table(self, k_results: List[KCalculationResult]) -> str:
        """Generate HTML table with K-value statistics: Method, K-value, Status, Applicability Range"""
        html = """
        <table class="table-compact" data-report-table="k-method-details">
            <tr>
                <th>Method</th>
                <th>K-Value (m/s)</th>
                <th>Status</th>
                <th>Applicability Notes</th>
            </tr>
        """

        for result in k_results:
            status_text = classify_k_status(result)
            k_display = f"{result.k_value:.2e}" if result.k_value else "N/A"
            notes = result.status_message if hasattr(result, 'status_message') and result.status_message else str(result.status)

            html += f"""
            <tr>
                <td><strong>{result.method_name}</strong></td>
                <td style="text-align: right;">{k_display}</td>
                <td style="text-align: center;">{status_text}</td>
                <td>{notes}</td>
            </tr>
            """

        html += "</table>"
        return html

    def _create_comparison_k_scope_summary_table(self, comparison_snapshot) -> str:
        """Render aggregate K statistics as two portrait-friendly tables."""
        k_report = comparison_snapshot.k
        grouped = self._comparison_uses_grouped_k_scope(comparison_snapshot)
        rows = [("Overall", k_report.overall)]
        if grouped:
            rows.extend(
                (group_name, k_report.by_group.get(group_name))
                for group_name in k_report.group_names
            )
        else:
            rows.extend(
                (dataset_name, k_report.by_dataset.get(dataset_name))
                for dataset_name in k_report.dataset_names
            )

        def fmt_k(value):
            return "N/A" if value is None else f"{value:.2e}"

        def fmt_float(value, precision=2):
            return "N/A" if value is None else f"{value:.{precision}f}"

        statistics = """
        <div class="table-pair">
        <table class="table-compact" data-report-table="k-aggregate-statistics">
            <thead><tr>
                <th>Scope</th><th class="num">Datasets</th>
                <th class="num">Geometric mean (m/s)</th>
                <th class="num">Arithmetic mean (m/s)</th>
                <th class="num">Median (m/s)</th>
                <th class="num">ln(K) std dev</th>
            </tr></thead><tbody>
        """
        coverage = """
        <table data-report-table="k-aggregate-coverage">
            <thead><tr>
                <th>Scope</th><th class="center">Included K cells</th>
                <th class="center">Warning cells</th><th>Permeability class</th>
            </tr></thead><tbody>
        """

        written = 0
        for scope_name, stats in rows:
            if stats is None:
                continue
            permeability = (
                _gc_perm_class(stats.geometric_mean_m_s)
                if stats.geometric_mean_m_s is not None
                else "N/A"
            )
            escaped_scope = self._esc(scope_name)
            statistics += f"""
                <tr>
                    <td><strong>{escaped_scope}</strong></td>
                    <td class="num">{stats.dataset_count}</td>
                    <td class="num">{fmt_k(stats.geometric_mean_m_s)}</td>
                    <td class="num">{fmt_k(stats.arithmetic_mean_m_s)}</td>
                    <td class="num">{fmt_k(stats.median_m_s)}</td>
                    <td class="num">{fmt_float(stats.ln_std_dev, 3)}</td>
                </tr>
            """
            coverage += f"""
                <tr>
                    <td><strong>{escaped_scope}</strong></td>
                    <td class="center">{stats.included_count} / {stats.total_cells}</td>
                    <td class="center">{stats.warning_count}</td>
                    <td>{self._esc(permeability)}</td>
                </tr>
            """
            written += 1

        if written == 0:
            statistics += '<tr><td colspan="6" class="table-empty">No included K-value results available</td></tr>'
            coverage += '<tr><td colspan="4" class="table-empty">No included K-value results available</td></tr>'

        return (
            statistics + "</tbody></table>"
            + coverage + "</tbody></table></div>"
        )

    def _create_comparison_k_statistics_table(self, k_results_dict: Dict[str, List[KCalculationResult]]) -> str:
        """Generate one compact row per sample/method result in canonical m/s."""
        html = """
        <table class="table-compact" data-report-table="k-method-results-comparison">
            <thead><tr>
                <th>Sample</th><th>Method</th><th class="num">K (m/s)</th>
                <th class="center">Status</th><th class="center">Included</th><th>Notes</th>
            </tr></thead><tbody>
        """

        row_count = 0
        for sample_name, results in k_results_dict.items():
            if not results:
                html += f"""
                <tr class="table-group-start">
                    <td><strong>{self._esc(sample_name)}</strong></td>
                    <td colspan="5" class="table-empty">No K-value results available</td>
                </tr>
                """
                continue

            for result_index, result in enumerate(results):
                status_text = classify_k_status(result)
                k_value = getattr(result, "k_value", None)
                has_value = k_value is not None and np.isfinite(k_value) and k_value > 0
                k_ms = f"{k_value:.2e}" if has_value else "N/A"
                included = has_value and status_text == "OK"
                notes = getattr(result, "status_message", "") or status_text
                row_class = ' class="table-group-start"' if result_index == 0 else ""

                html += f"""
                <tr{row_class}>
                    <td><strong>{self._esc(sample_name)}</strong></td>
                    <td>{self._esc(getattr(result, "method_name", ""))}</td>
                    <td class="num">{k_ms}</td>
                    <td class="center">{self._esc(status_text)}</td>
                    <td class="center">{"Yes" if included else "No"}</td>
                    <td class="note">{self._esc(notes)}</td>
                </tr>
                """
                row_count += 1

        if row_count == 0:
            html += '<tr><td colspan="6" class="table-empty">No K-value results available</td></tr>'

        return html + "</tbody></table>"

    def _create_data_quality_table(self, dataset: GrainSizeData) -> str:
        """Generate HTML table showing data quality metrics"""
        n_points = len(dataset.particle_sizes)
        size_min = min(dataset.particle_sizes)
        size_max = max(dataset.particle_sizes)
        size_range = size_max / size_min if size_min > 0 else 0

        # Check monotonicity
        sorted_indices = np.argsort(dataset.particle_sizes)[::-1]
        sorted_passing = [dataset.percent_passing[i] for i in sorted_indices]

        monotonic = all(sorted_passing[i] >= sorted_passing[i+1] for i in range(len(sorted_passing)-1))
        monotonicity_score = "Excellent" if monotonic else "Good"

        # Data coverage (log scale)
        coverage_score = "Excellent" if size_range > 100 else "Good" if size_range > 10 else "Limited"

        # Point density
        avg_spacing = np.mean([abs(dataset.particle_sizes[i] - dataset.particle_sizes[i-1])
                               for i in range(1, len(dataset.particle_sizes))])
        density_score = "Excellent" if n_points > 20 else "Good" if n_points > 10 else "Adequate"

        # Interpolation confidence
        confidence_score = "High" if (n_points > 15 and size_range > 50) else "Moderate" if n_points > 8 else "Low"

        html = """
        <table data-report-table="data-quality">
            <tr>
                <th>Quality Metric</th>
                <th>Value</th>
                <th>Assessment</th>
            </tr>
        """

        metrics = [
            ("Number of Data Points", str(n_points), density_score),
            ("Size Range", f"{size_min:.3f} - {size_max:.1f} mm", coverage_score),
            ("Span Ratio", f"{size_range:.1f}x", coverage_score),
            ("Curve Monotonicity", "Monotonic" if monotonic else "Some variation", monotonicity_score),
            ("Interpolation Confidence", "", confidence_score)
        ]

        for metric, value, assessment in metrics:
            # Color code assessment
            if assessment in ["Excellent", "High"]:
                color = "#e8f5e9"
            elif assessment in ["Good", "Moderate"]:
                color = "#fff9e6"
            else:
                color = "#ffebee"

            html += f"""
            <tr>
                <td style="font-weight: bold;">{metric}</td>
                <td>{value}</td>
                <td style="background-color: {color}; text-align: center;">{assessment}</td>
            </tr>
            """

        html += "</table>"
        return html

    def _create_sample_overview_tables(
        self,
        sample_details: List[Dict[str, Any]],
        mean_k_by_sample: Dict[str, float],
    ) -> str:
        """Render sample properties separately from classification/K results."""
        properties = """
        <div class="table-pair">
        <table class="table-compact" data-report-table="sample-properties">
            <thead><tr>
                <th>Sample</th><th class="num">Temp (?C)</th><th class="num">Porosity</th>
                <th class="num">D10 (mm)</th><th class="num">D50 (mm)</th>
                <th class="num">D60 (mm)</th><th class="num">Cu</th>
            </tr></thead><tbody>
        """
        classification = """
        <table data-report-table="sample-classification">
            <thead><tr>
                <th>Sample</th><th>Soil type</th><th>Descriptor</th><th class="num">K geometric mean (m/s)</th>
            </tr></thead><tbody>
        """

        for item in sample_details:
            label = str(item["label"])
            dataset = item["dataset"]
            d10 = dataset.get_d10()
            d50 = dataset.get_d50()
            d60 = dataset.get_d60()
            cu = (d60 / d10) if (d10 and d60) else None
            temp_value = item.get("temperature")
            porosity_value = item.get("porosity")
            temp_display = "N/A" if temp_value is None else f"{float(temp_value):.2f}"
            porosity_display = (
                "N/A" if porosity_value is None
                else f"{float(porosity_value):.3f}".rstrip("0").rstrip(".")
            )
            mean_k = mean_k_by_sample.get(label)
            mean_display = "N/A" if mean_k is None else f"{mean_k:.2e}"

            properties += f"""
                <tr>
                    <td>{self._esc(label)}</td>
                    <td class="num">{temp_display}</td>
                    <td class="num">{porosity_display}</td>
                    <td class="num">{f'{d10:.3f}' if d10 else 'N/A'}</td>
                    <td class="num">{f'{d50:.3f}' if d50 else 'N/A'}</td>
                    <td class="num">{f'{d60:.3f}' if d60 else 'N/A'}</td>
                    <td class="num">{f'{cu:.2f}' if cu else 'N/A'}</td>
                </tr>
            """
            _cls = dataset.classify(scheme=self._scheme)
            _descriptor = _gc_sed_descriptor(_cls.fractions, d50, cu, self._scheme)
            classification += f"""
                <tr>
                    <td>{self._esc(label)}</td>
                    <td>{self._esc(_cls.label)}</td>
                    <td>{self._esc(_descriptor or '—')}</td>
                    <td class="num">{mean_display}</td>
                </tr>
            """

        return (
            properties + "</tbody></table>"
            + classification + "</tbody></table></div>"
        )

    def _create_grain_parameters_comparison_table(self, datasets: List[GrainSizeData],
                                                  sample_labels: Optional[List[str]] = None) -> str:
        """Compare key grain parameters without creating an unprintable wide table."""
        labels = sample_labels or [dataset.sample_name for dataset in datasets]
        param_specs = [
            ("D10 (mm)", lambda ds: ds.get_d10(), ".3f"),
            ("D50 (mm)", lambda ds: ds.get_d50(), ".3f"),
            ("D60 (mm)", lambda ds: ds.get_d60(), ".3f"),
            ("Cu", lambda ds: ds.get_uniformity_coefficient(), ".2f"),
            ("Cc", lambda ds: ds.get_coefficient_of_curvature(), ".2f"),
        ]

        def fmt_value(value: Optional[float], fmt: str) -> str:
            return "-" if value is None else format(value, fmt)

        values_by_param = []
        for param, getter, fmt in param_specs:
            values = []
            for dataset in datasets:
                try:
                    values.append(getter(dataset))
                except Exception:
                    values.append(None)
            valid_values = [value for value in values if value is not None]
            if valid_values:
                mean = float(np.mean(valid_values))
                std = float(np.std(valid_values))
                summary = f"Mean: {format(mean, fmt)}<br>Std: {format(std, fmt)}"
            else:
                summary = "-"
            values_by_param.append((param, fmt, values, summary))

        if len(datasets) > 6:
            html = """
            <table class="table-compact" data-report-table="grain-parameters">
                <thead><tr><th>Parameter</th><th>Sample</th><th>Value</th></tr></thead>
                <tbody>
            """
            for param, fmt, values, summary in values_by_param:
                for row_index, (label, value) in enumerate(zip(labels, values)):
                    row_class = ' class="table-group-start"' if row_index == 0 else ""
                    html += f"""
                    <tr{row_class}>
                        <td><strong>{self._esc(param)}</strong></td>
                        <td>{self._esc(label)}</td>
                        <td class="num">{self._esc(fmt_value(value, fmt))}</td>
                    </tr>
                    """
                html += f"""
                <tr>
                    <td><strong>{self._esc(param)} summary</strong></td>
                    <td>Included samples</td>
                    <td class="num">{summary}</td>
                </tr>
                """
            return html + "</tbody></table>"

        html = """
        <table class="table-compact table-wide" data-report-table="grain-parameters">
            <thead><tr><th>Parameter</th>
        """
        for label in labels:
            html += f"<th class='num'>{self._esc(label)}</th>"
        html += "<th class='num'>Statistics</th></tr></thead><tbody>"

        for param, fmt, values, summary in values_by_param:
            html += f"<tr><td><strong>{self._esc(param)}</strong></td>"
            for value in values:
                html += f"<td class='num'>{self._esc(fmt_value(value, fmt))}</td>"
            html += f"<td class='num'>{summary}</td></tr>"

        return html + "</tbody></table>"

    def _create_permeability_classification_table(self, k_results_dict: Dict[str, List[KCalculationResult]]) -> str:
        """Generate HTML table with sample name, K geometric mean, and classification."""
        html = """
        <table data-report-table="permeability-summary">
            <thead>
                <tr>
                    <th>Sample Name</th>
                    <th>K geometric mean (m/s)</th>
                    <th>Classification</th>
                </tr>
            </thead>
            <tbody>
        """

        for sample_name, results in k_results_dict.items():
            summary = build_k_result_summary(results, dataset_name=sample_name)

            if summary.geometric_mean_m_s is not None:
                mean_k = summary.geometric_mean_m_s

                classification = _gc_perm_class(mean_k)

                html += f"""
                <tr>
                    <td><strong>{sample_name}</strong></td>
                    <td style="text-align: right;">{mean_k:.2e}</td>
                    <td>{classification}</td>
                </tr>
                """
            else:
                html += f"""
                <tr>
                    <td><strong>{sample_name}</strong></td>
                    <td style="text-align: center;">—</td>
                    <td style="text-align: center;">—</td>
                </tr>
                """

        html += "</tbody></table>"
        return html

    def _classify_permeability(self, k: float) -> str:
        return _gc_perm_class(k)

    def _get_permeability_application(self, k: float) -> str:
        if k > 1e-2:
            return "Excellent for drainage, unsuitable for water retention"
        elif k > 1e-4:
            return "Good for drainage systems, aquifers"
        elif k > 1e-5:
            return "Suitable for sand filters, moderate drainage"
        elif k > 1e-7:
            return "Poor drainage, may require improvement for construction"
        elif k > 1e-9:
            return "Natural barrier, suitable for liner with treatment"
        else:
            return "Excellent barrier material, natural aquitard"

    def _interpret_grain_distribution(self, dataset: GrainSizeData, cu: Optional[float], cc: Optional[float]) -> str:
        interpretation = f"The sample '{dataset.sample_name}' has been classified as {dataset.classify(scheme=self._scheme).label}. "

        if cu:
            cu_class = _gc_cu_label(cu)
            if cu_class == "Uniform":
                interpretation += "The uniform gradation (Cu < 4) indicates particles of similar size, "
                interpretation += "which typically results in higher void ratios and permeability. "
            elif cu_class == "Moderately graded":
                interpretation += "The moderate gradation (4 ≤ Cu < 6) suggests a reasonable distribution of particle sizes. "
            else:
                interpretation += "The well-graded nature (Cu ≥ 6) indicates a wide range of particle sizes, "
                interpretation += "which typically results in better compaction and lower permeability. "

        if cc and cu and cu >= 6:
            if 1 <= cc <= 3:
                interpretation += "The coefficient of curvature confirms well-graded material with good particle size distribution. "
            else:
                interpretation += "However, the coefficient of curvature suggests some gap-grading in the distribution. "

        return interpretation

    def _interpret_k_variability(self, ratio: float) -> str:
        if ratio < 10:
            return "The relatively low variability between methods suggests consistent and reliable results."
        elif ratio < 100:
            return "Moderate variability between methods is typical for this type of analysis. Consider using the median value."
        else:
            return "High variability between methods indicates uncertainty. Review input parameters and consider site-specific calibration."

    def _generate_no_results_report(self, sample_name: str) -> str:
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>K-Value Report - {sample_name}</title>
            {self.report_style}
        </head>
        <body>
            <h1>Hydraulic Conductivity Analysis Report</h1>
            <div class="warning-box">
                <h3>No Valid Results</h3>
                <p>No valid K-value calculations were obtained for sample '{sample_name}'.</p>
                <p>This may be due to:</p>
                <ul>
                    <li>Grain size parameters outside method applicability ranges</li>
                    <li>Missing required grain size data (D10, D60, etc.)</li>
                    <li>Invalid input parameters</li>
                </ul>
                <p>Please review the input data and ensure all required parameters are available.</p>
            </div>
        </body>
        </html>
        """
